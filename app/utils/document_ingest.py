from __future__ import annotations

import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

import fitz
import requests
from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[2]
DATA_ROOT = REPO_ROOT / "data"
OCR_ROOT = DATA_ROOT / "ocr"
OCR_TESSDATA_DIR = OCR_ROOT / "tessdata"
TESSERACT_URLS = {
    "chi_tra": "https://github.com/tesseract-ocr/tessdata_best/raw/main/chi_tra.traineddata",
    "eng": "https://github.com/tesseract-ocr/tessdata_best/raw/main/eng.traineddata",
}
TESSERACT_CANDIDATES = [
    Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe"),
    Path(r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"),
]


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def slugify(text: str) -> str:
    normalized = normalize_whitespace(text)
    normalized = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", normalized, flags=re.UNICODE)
    normalized = re.sub(r"-{2,}", "-", normalized).strip("-")
    return normalized.lower() or "item"


def normalize_whitespace(text: str) -> str:
    return " ".join(str(text or "").replace("\u3000", " ").replace("\xa0", " ").split())


def normalize_multiline_text(text: str) -> str:
    lines = [normalize_whitespace(line) for line in str(text or "").splitlines()]
    return "\n".join(line for line in lines if line)


def normalize_for_dedupe(text: str) -> str:
    text = normalize_multiline_text(text).lower()
    text = re.sub(r"\s+", "", text)
    return text


def sanitize_xml_text(text: str) -> str:
    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", str(text or ""))


def load_json(path: Path, default=None):
    if not path.exists():
        return {} if default is None else default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {} if default is None else default


def write_json(path: Path, payload) -> None:
    ensure_dir(path.parent)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def hash_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while True:
            chunk = handle.read(chunk_size)
            if not chunk:
                break
            digest.update(chunk)
    return digest.hexdigest()


def file_size_mb(path: Path) -> float:
    return round(path.stat().st_size / (1024 * 1024), 2)


def find_tesseract() -> Path | None:
    env_path = os.environ.get("TESSERACT_EXE")
    if env_path and Path(env_path).exists():
        return Path(env_path)

    which_path = shutil.which("tesseract")
    if which_path:
        return Path(which_path)

    for candidate in TESSERACT_CANDIDATES:
        if candidate.exists():
            return candidate
    return None


def ensure_ocr_language(language: str) -> Path:
    ensure_dir(OCR_TESSDATA_DIR)
    target = OCR_TESSDATA_DIR / f"{language}.traineddata"
    if target.exists():
        return target

    url = TESSERACT_URLS.get(language)
    if not url:
        raise RuntimeError(f"Missing OCR language source for {language}")

    response = requests.get(url, timeout=180)
    response.raise_for_status()
    target.write_bytes(response.content)
    return target


def ensure_ocr_assets(languages: str = "chi_tra+eng") -> dict[str, str]:
    tesseract_path = find_tesseract()
    if not tesseract_path:
        raise RuntimeError("Tesseract OCR is not installed on this machine.")

    for language in {token.strip() for token in languages.split("+") if token.strip()}:
        ensure_ocr_language(language)

    return {
        "tesseract_path": str(tesseract_path),
        "tessdata_dir": str(OCR_TESSDATA_DIR),
        "languages": languages,
    }


def score_text_quality(text: str) -> int:
    text = str(text or "")
    chinese_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    latin_count = len(re.findall(r"[A-Za-z]", text))
    question_like = len(re.findall(r"\b\d{1,2}[\.．、)]", text))
    return chinese_count + latin_count + question_like * 12


def run_tesseract_on_image(
    image_bytes: bytes,
    languages: str = "chi_tra+eng",
    psm: int = 6,
) -> str:
    assets = ensure_ocr_assets(languages)
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        image_path = tmp_path / "page.png"
        output_base = tmp_path / "ocr_output"
        image_path.write_bytes(image_bytes)

        command = [
            assets["tesseract_path"],
            str(image_path),
            str(output_base),
            "--tessdata-dir",
            assets["tessdata_dir"],
            "-l",
            assets["languages"],
            "--psm",
            str(psm),
        ]
        subprocess.run(command, check=True, capture_output=True)
        output_path = output_base.with_suffix(".txt")
        if not output_path.exists():
            return ""
        return output_path.read_text(encoding="utf-8", errors="ignore")


def render_page_png(page: fitz.Page, dpi: int = 220) -> bytes:
    pixmap = page.get_pixmap(dpi=dpi, alpha=False)
    return pixmap.tobytes("png")


def extract_pdf_pages(
    pdf_path: Path,
    languages: str = "chi_tra+eng",
    ocr_threshold: int = 80,
) -> list[dict]:
    pages = []
    with fitz.open(pdf_path) as document:
        for page_index in range(document.page_count):
            page = document[page_index]
            direct_text = normalize_multiline_text(page.get_text("text"))
            direct_score = score_text_quality(direct_text)
            final_text = direct_text
            source = "direct"
            ocr_text = ""

            if direct_score < ocr_threshold:
                try:
                    ocr_text = normalize_multiline_text(run_tesseract_on_image(render_page_png(page), languages=languages))
                except Exception:
                    ocr_text = ""
                if score_text_quality(ocr_text) > direct_score:
                    final_text = ocr_text
                    source = "ocr"
                elif ocr_text:
                    final_text = "\n".join(filter(None, [direct_text, ocr_text]))
                    source = "direct+ocr"

            pages.append(
                {
                    "page_number": page_index + 1,
                    "text": final_text,
                    "direct_text": direct_text,
                    "ocr_text": ocr_text,
                    "source": source,
                    "quality_score": score_text_quality(final_text),
                }
            )

    return pages


def save_extraction_docx(
    output_path: Path,
    title: str,
    pages: list[dict],
    metadata: dict | None = None,
) -> None:
    ensure_dir(output_path.parent)
    document = Document()
    document.add_heading(title, level=0)

    for key, value in (metadata or {}).items():
        if value:
            document.add_paragraph(f"{key}：{value}")

    for page in pages:
        document.add_heading(f"第 {page['page_number']} 頁", level=1)
        paragraph = document.add_paragraph()
        paragraph.add_run(sanitize_xml_text(page.get("text") or "（本頁無法抽取文字）"))

    document.save(output_path)


def download_to_path(url: str, destination: Path, session: requests.Session | None = None) -> Path:
    ensure_dir(destination.parent)
    if destination.exists():
        return destination
    client = session or requests.Session()
    response = client.get(url, timeout=240, stream=True)
    response.raise_for_status()
    with destination.open("wb") as handle:
        for chunk in response.iter_content(chunk_size=1024 * 1024):
            if chunk:
                handle.write(chunk)
    return destination


def extract_blocks_with_positions(pdf_path: Path) -> list[dict]:
    blocks = []
    with fitz.open(pdf_path) as document:
        for page_index in range(document.page_count):
            for block in document[page_index].get_text("blocks"):
                x0, y0, x1, y1, text, *_ = block
                cleaned = normalize_multiline_text(text)
                if not cleaned:
                    continue
                blocks.append(
                    {
                        "page_number": page_index + 1,
                        "x0": x0,
                        "y0": y0,
                        "x1": x1,
                        "y1": y1,
                        "text": cleaned,
                    }
                )
    return blocks


def read_words(pdf_path: Path, page_number: int) -> list[dict]:
    with fitz.open(pdf_path) as document:
        page = document[page_number - 1]
        words = []
        for item in page.get_text("words"):
            x0, y0, x1, y1, text, *_ = item
            if normalize_whitespace(text):
                words.append({"x0": x0, "y0": y0, "x1": x1, "y1": y1, "text": text})
        return words
