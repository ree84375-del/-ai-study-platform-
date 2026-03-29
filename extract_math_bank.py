from __future__ import annotations

import csv
import hashlib
import html
import logging
import re
import shutil
import time
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

import fitz
import requests
import urllib3
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from pdf2docx import Converter
from rapidocr_onnxruntime import RapidOCR


urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
logging.getLogger("pdf2docx").setLevel(logging.ERROR)
logging.getLogger("pdf2docx.converter").setLevel(logging.ERROR)


DOWNLOAD_PREFIX = "\u6578\u5b78"
QUESTION_SUFFIX = "_\u984c\u76ee\u5377"
ANSWER_SUFFIX = "_\u7b54\u6848\u5377"
OUTPUT_SUBDIR = "\u6578\u5b78"
READING_TXT = "\u6578\u5b78_\u95b1\u8b80\u7248.txt"
READING_DOCX = "\u6578\u5b78_\u95b1\u8b80\u7248.docx"
OUTPUT_CSV = "\u6578\u5b78_\u6240\u6709\u984c\u76ee.csv"
MISSING_REPORT = "\u6578\u5b78_\u7f3a\u4ef6\u6e05\u55ae.txt"
ISSUE_REPORT = "\u6578\u5b78_\u554f\u984c\u984c\u76ee\u6e05\u55ae.txt"
TMP_DIRNAME = "_tmp_math_extract"
TARGET_PATHS_FILE = "math_target_paths.txt"

CIRN_BASE_URL = "https://cirn.moe.edu.tw"
CIRN_VOLUME_URLS = {
    "\u7b2c\u4e00\u518a": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13216",
    "\u7b2c\u4e8c\u518a": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13217",
    "\u7b2c\u4e09\u518a": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13218",
    "\u7b2c\u56db\u518a": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13240",
    "\u7b2c\u4e94\u518a": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13241",
    "\u7b2c\u516d\u518a": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13242",
}

ROW_RE = re.compile(
    r'<span id="ctl00_ContentPlaceHolder1_gv1_ctl(?P<row>\d+)_lbFileName">(?P<name>[^<]+)</span>.*?'
    r"javascript:__doPostBack\(&#39;(?P<target>[^&#]+)&#39;,&#39;&#39;\)",
    re.S,
)
QUESTION_START_PATTERNS = [
    re.compile(r"^[\uff08(]\s*[\uff09)]\s*(\d+)\s*[\u3001.:\uff1a]?\s*(.*)$"),
    re.compile(r"^\u554f\u984c\s*(\d+)\s*[\u3001.:\uff1a]\s*(.*)$"),
]
ANSWER_START_PATTERNS = [
    re.compile(r"^[\uff08(]\s*([A-D])\s*[\uff09)]\s*(\d+)\s*[\u3001.:\uff1a]?\s*(.*)$", re.I),
    re.compile(r"^[\uff08(]\s*([A-D])\s*[\uff09)]\s*[\uff08(]?\s*(\d+)\s*[\uff09)]?\s*[\u3001.:\uff1a]?\s*(.*)$", re.I),
    re.compile(r"^([A-D])\s*[\u3001.:\uff1a)]\s*(\d+)\s*[\u3001.:\uff1a]?\s*(.*)$", re.I),
]
OPTION_START_PATTERNS = [
    re.compile(r"^[\uff08(]?\s*([A-D])\s*[\uff09).:\uff1a]\s*(.*)$", re.I),
]
BOOKLET_RE = re.compile(r"^\u6578\u5b78(\u7b2c[一二三四五六1-6]\u518a)_")
SUBITEM_RE = re.compile(r"^[\uff08(]?\s*\d+\s*[\uff09)]")

OCR_ENGINE = RapidOCR()


@dataclass
class SourcePair:
    unit_key: str
    question_candidates: list[Path] = field(default_factory=list)
    answer_candidates: list[Path] = field(default_factory=list)
    question_pdf: Path | None = None
    answer_pdf: Path | None = None


@dataclass
class QuestionBlock:
    number: int
    answer: str = ""
    lines: list[str] = field(default_factory=list)


@dataclass
class CandidateParts:
    question_text: str
    options: list[str]
    source_name: str
    score: int


@dataclass
class MathQuestion:
    booklet: str
    chapter: str
    topic: str
    source_unit: str
    question_number: int
    question_text: str
    option_a: str
    option_b: str
    option_c: str
    option_d: str
    correct_answer: str
    explanation: str


@dataclass
class IssueItem:
    kind: str
    unit_key: str
    question_number: int | None
    detail: str


def get_downloads_dir() -> Path:
    return Path.home() / "Downloads"


def get_output_dir() -> Path:
    output_dir = Path.home() / "Desktop" / "\u570b\u4e2d\u984c\u5eab" / OUTPUT_SUBDIR
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def get_target_paths_file() -> Path:
    return Path.cwd() / TARGET_PATHS_FILE


def normalize_pdf_name(name: str) -> str:
    cleaned = re.sub(rf"{re.escape(QUESTION_SUFFIX)}\s*\(\d+\)(?=\.pdf$)", QUESTION_SUFFIX, name)
    cleaned = re.sub(rf"{re.escape(ANSWER_SUFFIX)}\s*\(\d+\)(?=\.pdf$)", ANSWER_SUFFIX, cleaned)
    cleaned = re.sub(r" \(\d+\)(?=\.pdf$)", "", cleaned)
    return cleaned


def unit_key_from_filename(name: str) -> tuple[str, str | None]:
    if QUESTION_SUFFIX in name:
        return name.split(QUESTION_SUFFIX)[0], "question"
    if ANSWER_SUFFIX in name:
        return name.split(ANSWER_SUFFIX)[0], "answer"
    return name.rsplit(".", 1)[0], None


def duplicate_penalty(path: Path) -> tuple[int, int, str]:
    normalized = normalize_pdf_name(path.name)
    return (
        1 if path.name != normalized else 0,
        len(path.name),
        path.name,
    )


def choose_preferred(paths: list[Path]) -> Path | None:
    if not paths:
        return None
    return sorted(paths, key=duplicate_penalty)[0]


def load_target_unit_keys(target_paths_file: Path) -> set[str]:
    unit_keys: set[str] = set()
    if not target_paths_file.exists():
        return unit_keys

    for raw_line in target_paths_file.read_text(encoding="utf-8-sig").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        normalized_name = normalize_pdf_name(Path(line).name)
        unit_key, kind = unit_key_from_filename(normalized_name)
        if kind:
            unit_keys.add(unit_key)
    return unit_keys


def collect_input_paths(downloads_dir: Path, target_unit_keys: set[str]) -> list[Path]:
    pdf_paths = [path for path in downloads_dir.iterdir() if path.suffix.lower() == ".pdf" and path.name.startswith(DOWNLOAD_PREFIX)]
    if not target_unit_keys:
        return pdf_paths

    filtered_paths: list[Path] = []
    for path in pdf_paths:
        normalized_name = normalize_pdf_name(path.name)
        unit_key, kind = unit_key_from_filename(normalized_name)
        if kind and unit_key in target_unit_keys:
            filtered_paths.append(path)
    return filtered_paths


def collect_pairs(paths: Iterable[Path]) -> dict[str, SourcePair]:
    pairs: dict[str, SourcePair] = {}
    for path in sorted(paths, key=lambda item: item.name):
        normalized_name = normalize_pdf_name(path.name)
        unit_key, kind = unit_key_from_filename(normalized_name)
        pair = pairs.setdefault(unit_key, SourcePair(unit_key=unit_key))
        if kind == "question":
            pair.question_candidates.append(path)
        elif kind == "answer":
            pair.answer_candidates.append(path)
    for pair in pairs.values():
        pair.question_pdf = choose_preferred(pair.question_candidates)
        pair.answer_pdf = choose_preferred(pair.answer_candidates)
    return pairs


def build_hidden_field_map(page: str) -> dict[str, str]:
    return {
        match.group(1): html.unescape(match.group(2))
        for match in re.finditer(
            r'<input[^>]+type="hidden"[^>]+name="([^"]+)"[^>]+value="([^"]*)"',
            page,
        )
    }


def fetch_all_rows_for_volume(url: str) -> tuple[requests.Session, str]:
    session = requests.Session()
    page = session.get(url, timeout=30, verify=False).text
    fields = build_hidden_field_map(page)
    fields["ctl00$ContentPlaceHolder1$PagingControl1$txtPageCut"] = "200"
    fields["ctl00$ContentPlaceHolder1$PagingControl1$btnNoUsed"] = ""
    page = session.post(url, data=fields, timeout=30, verify=False).text
    return session, page


def download_file_from_cirn(
    session: requests.Session,
    page: str,
    url: str,
    event_target: str,
    output_path: Path,
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    fields = build_hidden_field_map(page)
    fields["ctl00$ContentPlaceHolder1$PagingControl1$txtPageCut"] = "200"
    fields["__EVENTTARGET"] = event_target
    fields["__EVENTARGUMENT"] = ""
    response = session.post(url, data=fields, timeout=60, verify=False)
    response.raise_for_status()
    output_path.write_bytes(response.content)
    return output_path


def volume_from_unit_key(unit_key: str) -> str:
    match = BOOKLET_RE.match(unit_key)
    if not match:
        raise ValueError(f"Cannot parse booklet from {unit_key}")
    booklet = match.group(1)
    if booklet.endswith("1\u518a"):
        return "\u7b2c\u4e00\u518a"
    if booklet.endswith("2\u518a"):
        return "\u7b2c\u4e8c\u518a"
    if booklet.endswith("3\u518a"):
        return "\u7b2c\u4e09\u518a"
    if booklet.endswith("4\u518a"):
        return "\u7b2c\u56db\u518a"
    if booklet.endswith("5\u518a"):
        return "\u7b2c\u4e94\u518a"
    if booklet.endswith("6\u518a"):
        return "\u7b2c\u516d\u518a"
    return booklet


def download_missing_from_cirn(pairs: dict[str, SourcePair], temp_dir: Path) -> list[str]:
    downloaded: list[str] = []
    for booklet, url in CIRN_VOLUME_URLS.items():
        volume_units = [
            pair
            for pair in pairs.values()
            if volume_from_unit_key(pair.unit_key) == booklet
            and (pair.question_pdf is None or pair.answer_pdf is None)
        ]
        if not volume_units:
            continue

        session, page = fetch_all_rows_for_volume(url)
        available_rows = {
            html.unescape(match.group("name")): html.unescape(match.group("target"))
            for match in ROW_RE.finditer(page)
        }

        for pair in volume_units:
            if pair.question_pdf is None:
                filename = f"{pair.unit_key}{QUESTION_SUFFIX}.pdf"
                target = available_rows.get(filename)
                if target:
                    downloaded_path = download_file_from_cirn(
                        session=session,
                        page=page,
                        url=url,
                        event_target=target,
                        output_path=temp_dir / filename,
                    )
                    pair.question_candidates.append(downloaded_path)
                    pair.question_pdf = downloaded_path
                    downloaded.append(filename)
            if pair.answer_pdf is None:
                filename = f"{pair.unit_key}{ANSWER_SUFFIX}.pdf"
                target = available_rows.get(filename)
                if target:
                    downloaded_path = download_file_from_cirn(
                        session=session,
                        page=page,
                        url=url,
                        event_target=target,
                        output_path=temp_dir / filename,
                    )
                    pair.answer_candidates.append(downloaded_path)
                    pair.answer_pdf = downloaded_path
                    downloaded.append(filename)
    return downloaded


def normalize_line(line: str) -> str:
    line = (
        line.replace("\xa0", " ")
        .replace("\u3000", " ")
        .replace("\u200b", "")
        .replace("\ufeff", "")
    )
    return re.sub(r"\s+", " ", line).strip()


def is_header_line(line: str) -> bool:
    if not line:
        return True
    if line.startswith("\u6578\u5b78\uff0f\u7b2c") or line.startswith("\u6578\u5b78/\u7b2c"):
        return True
    if line in {"\u3010\u984c\u76ee\u5377\u3011", "\u3010\u7b54\u6848\u5377\u3011"}:
        return True
    if line.startswith("\u9801") and re.search(r"\d+$", line):
        return True
    if re.fullmatch(r"\d+", line):
        return True
    return False


def read_pdf_lines(pdf_path: Path) -> list[str]:
    document = fitz.open(pdf_path)
    lines: list[str] = []
    for page in document:
        for raw_line in page.get_text("text").splitlines():
            line = normalize_line(raw_line)
            if not line or is_header_line(line):
                continue
            lines.append(line)
    document.close()
    return lines


def convert_pdf_to_docx(pdf_path: Path, output_path: Path) -> None:
    converter = Converter(str(pdf_path))
    try:
        converter.convert(str(output_path), start=0, end=None)
    finally:
        converter.close()


def iter_block_items(parent: DocumentObject | _Cell):
    parent_element = parent.element.body if isinstance(parent, DocumentObject) else parent._tc
    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def read_docx_lines(docx_path: Path) -> list[str]:
    document = Document(docx_path)
    lines: list[str] = []

    def append_line(value: str) -> None:
        line = normalize_line(value)
        if line and not is_header_line(line):
            lines.append(line)

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            append_line(block.text)
            continue
        for row in block.rows:
            for cell in row.cells:
                for item in iter_block_items(cell):
                    if isinstance(item, Paragraph):
                        append_line(item.text)
    return lines


def render_ocr_lines(pdf_path: Path, image_dir: Path) -> list[str]:
    document = fitz.open(pdf_path)
    lines: list[str] = []
    for index, page in enumerate(document):
        image_path = image_dir / f"{pdf_path.stem}_page_{index + 1}.png"
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2.2, 2.2), alpha=False)
        pixmap.save(image_path)
        ocr_result, _ = OCR_ENGINE(str(image_path))
        for item in ocr_result or []:
            line = normalize_line(str(item[1]))
            if line and not is_header_line(line):
                lines.append(line)
    document.close()
    return lines


def parse_blocks(lines: list[str], with_answers: bool) -> dict[int, QuestionBlock]:
    patterns = ANSWER_START_PATTERNS if with_answers else QUESTION_START_PATTERNS
    blocks: dict[int, QuestionBlock] = {}
    current: QuestionBlock | None = None
    preface_lines: list[str] = []

    def should_buffer_for_next_question(current_lines: list[str], incoming_line: str) -> bool:
        if with_answers:
            return False
        labeled_count = 0
        for existing_line in current_lines:
            if any(pattern.match(existing_line) for pattern in OPTION_START_PATTERNS):
                labeled_count += 1
        if labeled_count >= 4 and len(incoming_line) > 18 and not SUBITEM_RE.match(incoming_line):
            return True

        trailing_short = 0
        for existing_line in reversed(current_lines):
            if any(pattern.match(existing_line) for pattern in OPTION_START_PATTERNS):
                break
            if len(existing_line) <= 120:
                trailing_short += 1
            else:
                break
        if trailing_short >= 4 and len(incoming_line) > 18:
            return True
        return False

    for raw_line in lines:
        line = normalize_line(raw_line)
        if not line or is_header_line(line):
            continue

        matched = False
        for pattern in patterns:
            match = pattern.match(line)
            if not match:
                continue

            if current is not None:
                blocks[current.number] = current

            if with_answers:
                answer, number, rest = match.groups()
                current = QuestionBlock(number=int(number), answer=answer.upper())
            else:
                number, rest = match.groups()
                current = QuestionBlock(number=int(number))

            if preface_lines:
                current.lines.extend(preface_lines)
                preface_lines = []

            rest = normalize_line(rest)
            if rest:
                current.lines.append(rest)
            matched = True
            break

        if matched:
            continue

        if current is not None:
            if should_buffer_for_next_question(current.lines, line):
                preface_lines.append(line)
                continue
            current.lines.append(line)
        else:
            preface_lines.append(line)

    if current is not None:
        blocks[current.number] = current

    return blocks


def parse_labeled_options(lines: list[str]) -> tuple[list[str], list[str]]:
    question_lines: list[str] = []
    options_map: dict[str, list[str]] = {}
    current_option: str | None = None

    for line in lines:
        matched_option = False
        for pattern in OPTION_START_PATTERNS:
            match = pattern.match(line)
            if not match:
                continue
            label, rest = match.groups()
            label = label.upper()
            if label not in {"A", "B", "C", "D"}:
                continue
            options_map.setdefault(label, [])
            if rest:
                options_map[label].append(normalize_line(rest))
            current_option = label
            matched_option = True
            break

        if matched_option:
            continue

        if current_option is None:
            question_lines.append(line)
        else:
            options_map.setdefault(current_option, []).append(line)

    if set(options_map) >= {"A", "B", "C", "D"}:
        options = [normalize_line(" ".join(options_map.get(label, []))) for label in ("A", "B", "C", "D")]
        return question_lines, options
    return question_lines, []


def parse_trailing_options(lines: list[str]) -> tuple[str, list[str]]:
    if len(lines) < 5:
        return normalize_line(" ".join(lines)), []

    candidate = lines[-4:]
    if not all(line and len(line) <= 120 for line in candidate):
        return normalize_line(" ".join(lines)), []

    question_text = normalize_line(" ".join(lines[:-4]))
    options = [normalize_line(line) for line in candidate]
    return question_text, options


def score_candidate(question_text: str, options: list[str], source_name: str) -> int:
    score = min(len(question_text), 240)
    if len(options) == 4:
        score += 120
    elif len(options) == 3:
        score += 55
    else:
        score -= 80

    if source_name.endswith("_ocr"):
        score += 18
    if source_name.endswith("_direct"):
        score += 10
    if source_name.endswith("_docx"):
        score += 8

    repeated_option_markers = sum(question_text.count(token) for token in ("(1)", "\uff081\uff09", "(A)", "\uff08A\uff09"))
    if repeated_option_markers >= 2:
        score -= 150

    if len(question_text) < 8:
        score -= 80

    if any(not option for option in options):
        score -= 40

    return score


def build_candidate_parts(lines: list[str], source_name: str) -> CandidateParts | None:
    cleaned = [normalize_line(line) for line in lines if normalize_line(line)]
    if not cleaned:
        return None

    question_lines, labeled_options = parse_labeled_options(cleaned)
    if labeled_options:
        question_text = normalize_line(" ".join(question_lines))
        return CandidateParts(
            question_text=question_text,
            options=labeled_options,
            source_name=source_name,
            score=score_candidate(question_text, labeled_options, source_name),
        )

    question_text, trailing_options = parse_trailing_options(cleaned)
    return CandidateParts(
        question_text=question_text,
        options=trailing_options,
        source_name=source_name,
        score=score_candidate(question_text, trailing_options, source_name),
    )


def choose_best_parts(candidates: list[CandidateParts]) -> CandidateParts | None:
    if not candidates:
        return None
    return max(candidates, key=lambda item: item.score)


def extract_explanation(lines: list[str], answer: str) -> str:
    for index, line in enumerate(lines):
        if "\u89e3\u6790" in line or "\u8a73\u89e3" in line:
            explanation_lines = [line]
            explanation_lines.extend(lines[index + 1 :])
            explanation = normalize_line(" ".join(explanation_lines))
            if explanation:
                return explanation
    if answer:
        return f"\u6b63\u78ba\u7b54\u6848\u70ba {answer}\u3002\u539f\u59cb\u4f86\u6e90\u672a\u9644\u8a73\u89e3\uff0c\u5efa\u8b70\u5148\u56de\u982d\u6bd4\u5c0d\u984c\u5e79\u689d\u4ef6\u8207\u6b63\u78ba\u9078\u9805\u3002"
    return ""


def compute_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(65536), b""):
            digest.update(chunk)
    return digest.hexdigest()


def parse_unit_parts(unit_key: str) -> tuple[str, str, str]:
    parts = unit_key.split("_")
    booklet = parts[0] if parts else ""
    chapter = parts[1] if len(parts) > 1 else ""
    topic = "_".join(parts[2:]) if len(parts) > 2 else chapter
    return booklet, chapter, topic


def build_questions_for_pair(
    pair: SourcePair,
    temp_dir: Path,
    issues: list[IssueItem],
) -> list[MathQuestion]:
    if pair.question_pdf is None or pair.answer_pdf is None:
        return []

    question_direct_lines = read_pdf_lines(pair.question_pdf)
    answer_direct_lines = read_pdf_lines(pair.answer_pdf)

    output_dir = get_output_dir()
    cached_question_docx = output_dir / f"{pair.question_pdf.stem}.docx"
    question_docx_path = cached_question_docx if cached_question_docx.exists() else temp_dir / "docx" / f"{pair.question_pdf.stem}.docx"
    question_docx_path.parent.mkdir(parents=True, exist_ok=True)

    question_docx_lines: list[str] = []
    answer_docx_lines: list[str] = []

    try:
        if not question_docx_path.exists():
            convert_pdf_to_docx(pair.question_pdf, question_docx_path)
        question_docx_lines = read_docx_lines(question_docx_path)
    except Exception as exc:
        issues.append(IssueItem("docx_conversion_failed", pair.unit_key, None, f"question_pdf: {exc}"))

    question_blocks_direct = parse_blocks(question_direct_lines, with_answers=False)
    answer_blocks_direct = parse_blocks(answer_direct_lines, with_answers=True)
    question_blocks_docx = parse_blocks(question_docx_lines, with_answers=False)
    answer_blocks_docx = parse_blocks(answer_docx_lines, with_answers=True)

    question_blocks_ocr: dict[int, QuestionBlock] = {}
    answer_blocks_ocr: dict[int, QuestionBlock] = {}
    image_dir = temp_dir / "ocr_images"
    image_dir.mkdir(parents=True, exist_ok=True)

    question_numbers = set(question_blocks_direct) | set(question_blocks_docx) | set(answer_blocks_direct)

    low_quality_seed_count = 0
    for number in sorted(question_numbers):
        seed_candidates: list[CandidateParts] = []
        for source_name, block in [
            ("question_direct", question_blocks_direct.get(number)),
            ("question_docx", question_blocks_docx.get(number)),
            ("answer_direct", answer_blocks_direct.get(number)),
        ]:
            if not block:
                continue
            candidate = build_candidate_parts(block.lines, source_name=source_name)
            if candidate:
                seed_candidates.append(candidate)
        best_seed = choose_best_parts(seed_candidates)
        if best_seed is None or len(best_seed.question_text) < 8 or len([opt for opt in best_seed.options if opt]) < 4:
            low_quality_seed_count += 1

    need_question_ocr = len(question_blocks_direct) < len(answer_blocks_direct) or low_quality_seed_count > 0
    need_answer_ocr = len(answer_blocks_direct) == 0

    if need_question_ocr:
        question_ocr_lines = render_ocr_lines(pair.question_pdf, image_dir=image_dir)
        question_blocks_ocr = parse_blocks(question_ocr_lines, with_answers=False)
    if need_answer_ocr:
        answer_ocr_lines = render_ocr_lines(pair.answer_pdf, image_dir=image_dir)
        answer_blocks_ocr = parse_blocks(answer_ocr_lines, with_answers=True)

    question_numbers = set(question_blocks_direct) | set(question_blocks_docx) | set(question_blocks_ocr)
    answer_numbers = set(answer_blocks_direct) | set(answer_blocks_docx) | set(answer_blocks_ocr)

    missing_in_question = sorted(answer_numbers - question_numbers)
    missing_in_answer = sorted(question_numbers - answer_numbers)
    for number in missing_in_question:
        issues.append(IssueItem("answer_without_question", pair.unit_key, number, "\u7b54\u6848\u5377\u6709\u984c\u865f\uff0c\u4f46\u984c\u76ee\u5377\u6c92\u6293\u5230\u5c0d\u61c9\u984c\u76ee\u3002"))
    for number in missing_in_answer:
        issues.append(IssueItem("question_without_answer", pair.unit_key, number, "\u984c\u76ee\u5377\u6709\u984c\u865f\uff0c\u4f46\u7b54\u6848\u5377\u6c92\u6293\u5230\u5c0d\u61c9\u7b54\u6848\u3002"))

    booklet, chapter, topic = parse_unit_parts(pair.unit_key)
    questions: list[MathQuestion] = []

    for number in sorted(question_numbers | answer_numbers):
        block_candidates: list[CandidateParts] = []
        for source_name, block in [
            ("question_direct", question_blocks_direct.get(number)),
            ("question_docx", question_blocks_docx.get(number)),
            ("question_ocr", question_blocks_ocr.get(number)),
            ("answer_direct", answer_blocks_direct.get(number)),
            ("answer_docx", answer_blocks_docx.get(number)),
            ("answer_ocr", answer_blocks_ocr.get(number)),
        ]:
            if not block:
                continue
            candidate = build_candidate_parts(block.lines, source_name=source_name)
            if candidate:
                block_candidates.append(candidate)

        best_parts = choose_best_parts(block_candidates)
        if best_parts is None:
            issues.append(IssueItem("unparsed_question", pair.unit_key, number, "\u7121\u6cd5\u5f9e PDF / Word / OCR \u7d50\u679c\u91cd\u5efa\u984c\u76ee\u3002"))
            continue

        answers_found = [
            block.answer
            for block in (answer_blocks_direct.get(number), answer_blocks_docx.get(number), answer_blocks_ocr.get(number))
            if block and block.answer
        ]
        answer = Counter(answers_found).most_common(1)[0][0] if answers_found else ""
        if len(set(answers_found)) > 1:
            issues.append(
                IssueItem(
                    "answer_mismatch",
                    pair.unit_key,
                    number,
                    f"\u7b2c 1 \u8f2a\u548c\u7b2c 2 \u8f2a\u6aa2\u67e5\u51fa\u73fe\u4e0d\u540c\u7b54\u6848\uff1a{', '.join(sorted(set(answers_found)))}",
                )
            )

        option_count = len([option for option in best_parts.options if option])
        if len(best_parts.question_text) < 8 or option_count < 4 or answer not in {"A", "B", "C", "D"}:
            issues.append(
                IssueItem(
                    "low_quality_question",
                    pair.unit_key,
                    number,
                    f"question_len={len(best_parts.question_text)} option_count={option_count} answer={answer or 'EMPTY'} source={best_parts.source_name}",
                )
            )

        answer_block = answer_blocks_direct.get(number) or answer_blocks_docx.get(number) or answer_blocks_ocr.get(number)
        explanation_lines = answer_block.lines if answer_block else []
        explanation = extract_explanation(explanation_lines, answer=answer)

        options = (best_parts.options + ["", "", "", ""])[:4]
        questions.append(
            MathQuestion(
                booklet=booklet,
                chapter=chapter,
                topic=topic,
                source_unit=pair.unit_key,
                question_number=number,
                question_text=best_parts.question_text,
                option_a=options[0],
                option_b=options[1],
                option_c=options[2],
                option_d=options[3],
                correct_answer=answer,
                explanation=explanation,
            )
        )

    return questions


def normalized_question_key(question: MathQuestion) -> str:
    return normalize_line(question.question_text)


def deduplicate_questions(questions: list[MathQuestion], issues: list[IssueItem]) -> list[MathQuestion]:
    kept: list[MathQuestion] = []
    seen: dict[str, MathQuestion] = {}
    for question in questions:
        key = normalized_question_key(question)
        if not key:
            issues.append(IssueItem("empty_question", question.source_unit, question.question_number, "\u984c\u5e79\u70ba\u7a7a\uff0c\u5df2\u7565\u904e\u3002"))
            continue
        if key in seen:
            issues.append(
                IssueItem(
                    "duplicate_question_text",
                    question.source_unit,
                    question.question_number,
                    f"\u8207 {seen[key].source_unit} \u7b2c {seen[key].question_number} \u984c\u984c\u5e79\u5b8c\u5168\u76f8\u540c\uff0c\u5df2\u79fb\u9664\u91cd\u8907\u984c\u3002",
                )
            )
            continue
        seen[key] = question
        kept.append(question)
    return kept


def write_missing_report(path: Path, pairs: dict[str, SourcePair]) -> None:
    lines = ["\u6578\u5b78\u984c\u5eab\u7f3a\u4ef6\u6e05\u55ae", "=" * 32, ""]
    missing_rows = []
    for pair in sorted(pairs.values(), key=lambda item: item.unit_key):
        if pair.question_pdf is None:
            missing_rows.append((pair.unit_key, "\u7f3a\u984c\u76ee\u5377"))
        if pair.answer_pdf is None:
            missing_rows.append((pair.unit_key, "\u7f3a\u7b54\u6848\u5377"))

    if not missing_rows:
        lines.append("\u7121\u7f3a\u4ef6")
    else:
        for unit_key, detail in missing_rows:
            lines.append(f"- {unit_key}: {detail}")

    path.write_text("\n".join(lines), encoding="utf-8-sig")


def write_issue_report(path: Path, issues: list[IssueItem], pairs: dict[str, SourcePair]) -> None:
    sections: list[str] = ["\u6578\u5b78\u984c\u5eab\u554f\u984c\u984c\u76ee\u6e05\u55ae", "=" * 32, ""]

    duplicate_rows = []
    for pair in sorted(pairs.values(), key=lambda item: item.unit_key):
        for kind, candidates in (("\u984c\u76ee\u5377", pair.question_candidates), ("\u7b54\u6848\u5377", pair.answer_candidates)):
            if len(candidates) <= 1:
                continue
            hashes = {candidate.name: compute_sha256(candidate)[:12] for candidate in candidates}
            duplicate_rows.append((pair.unit_key, kind, hashes))

    sections.append("[\u91cd\u8907\u4e0b\u8f09]")
    if not duplicate_rows:
        sections.append("- \u7121")
    else:
        for unit_key, kind, hashes in duplicate_rows:
            sections.append(f"- {unit_key} / {kind}")
            for filename, short_hash in hashes.items():
                sections.append(f"  {filename} / sha256={short_hash}")
    sections.append("")

    grouped: dict[str, list[IssueItem]] = {}
    for issue in issues:
        grouped.setdefault(issue.kind, []).append(issue)

    for kind in sorted(grouped):
        sections.append(f"[{kind}]")
        for issue in grouped[kind]:
            if issue.question_number is None:
                sections.append(f"- {issue.unit_key}: {issue.detail}")
            else:
                sections.append(f"- {issue.unit_key} / \u7b2c {issue.question_number} \u984c: {issue.detail}")
        sections.append("")

    path.write_text("\n".join(sections), encoding="utf-8-sig")


def write_csv(path: Path, questions: list[MathQuestion]) -> None:
    fieldnames = [
        "volume",
        "category",
        "title",
        "source_unit",
        "question_no",
        "content_text",
        "option_a",
        "option_b",
        "option_c",
        "option_d",
        "correct_answer",
        "explanation",
        "tags",
        "difficulty",
        "booklet",
        "chapter",
        "topic",
    ]
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for question in questions:
            writer.writerow(
                {
                    "volume": question.booklet,
                    "category": question.chapter,
                    "title": question.topic,
                    "source_unit": question.source_unit,
                    "question_no": question.question_number,
                    "content_text": question.question_text,
                    "option_a": question.option_a,
                    "option_b": question.option_b,
                    "option_c": question.option_c,
                    "option_d": question.option_d,
                    "correct_answer": question.correct_answer,
                    "explanation": question.explanation,
                    "tags": f"{question.topic} | {question.source_unit}",
                    "difficulty": 2,
                    "booklet": question.booklet,
                    "chapter": question.chapter,
                    "topic": question.topic,
                }
            )


def write_reading_text(path: Path, questions: list[MathQuestion]) -> None:
    lines = ["\u6578\u5b78\u984c\u5eab\u95b1\u8b80\u7248", "=" * 32, ""]
    current_unit = None
    for question in questions:
        unit_marker = (question.booklet, question.chapter, question.topic)
        if unit_marker != current_unit:
            current_unit = unit_marker
            lines.extend(
                [
                    f"\u518a\u5225\uff1a{question.booklet}",
                    f"\u55ae\u5143\uff1a{question.chapter}",
                    f"\u4e3b\u984c\uff1a{question.topic}",
                    "",
                ]
            )
        lines.extend(
            [
                f"\u7b2c {question.question_number} \u984c",
                f"\u984c\u76ee\uff1a{question.question_text}",
                f"A. {question.option_a}",
                f"B. {question.option_b}",
                f"C. {question.option_c}",
                f"D. {question.option_d}",
                f"\u7b54\u6848\uff1a{question.correct_answer}",
                f"\u89e3\u6790\uff1a{question.explanation}",
                "",
            ]
        )
    path.write_text("\n".join(lines), encoding="utf-8-sig")


def write_reading_docx(path: Path, questions: list[MathQuestion]) -> None:
    document = Document()
    document.add_heading("\u6578\u5b78\u984c\u5eab\u95b1\u8b80\u7248", level=0)
    current_unit = None
    for question in questions:
        unit_marker = (question.booklet, question.chapter, question.topic)
        if unit_marker != current_unit:
            current_unit = unit_marker
            document.add_heading(f"{question.booklet} / {question.chapter} / {question.topic}", level=1)

        document.add_paragraph(f"\u7b2c {question.question_number} \u984c")
        document.add_paragraph(question.question_text)
        document.add_paragraph(f"A. {question.option_a}")
        document.add_paragraph(f"B. {question.option_b}")
        document.add_paragraph(f"C. {question.option_c}")
        document.add_paragraph(f"D. {question.option_d}")
        document.add_paragraph(f"\u7b54\u6848\uff1a{question.correct_answer}")
        document.add_paragraph(f"\u89e3\u6790\uff1a{question.explanation}")

    document.save(path)


def cleanup_old_temp_dirs(output_dir: Path) -> None:
    for child in output_dir.iterdir():
        if child.is_dir() and child.name.startswith(TMP_DIRNAME):
            shutil.rmtree(child, ignore_errors=True)


def main() -> None:
    downloads_dir = get_downloads_dir()
    output_dir = get_output_dir()
    target_paths_file = get_target_paths_file()
    target_unit_keys = load_target_unit_keys(target_paths_file)
    cleanup_old_temp_dirs(output_dir)
    temp_dir = output_dir / f"{TMP_DIRNAME}_{int(time.time())}"
    temp_dir.mkdir(parents=True, exist_ok=True)

    pdf_paths = collect_input_paths(downloads_dir, target_unit_keys)
    pairs = collect_pairs(pdf_paths)
    for unit_key in sorted(target_unit_keys):
        pairs.setdefault(unit_key, SourcePair(unit_key=unit_key))
    downloaded = download_missing_from_cirn(pairs, temp_dir=temp_dir / "downloads")

    issues: list[IssueItem] = []
    questions: list[MathQuestion] = []
    for pair in sorted(pairs.values(), key=lambda item: item.unit_key):
        if pair.question_pdf is None or pair.answer_pdf is None:
            continue
        questions.extend(build_questions_for_pair(pair, temp_dir=temp_dir, issues=issues))

    questions = deduplicate_questions(questions, issues=issues)
    questions.sort(key=lambda item: (item.booklet, item.chapter, item.topic, item.question_number))

    write_csv(output_dir / OUTPUT_CSV, questions)
    write_reading_text(output_dir / READING_TXT, questions)
    write_reading_docx(output_dir / READING_DOCX, questions)
    write_missing_report(output_dir / MISSING_REPORT, pairs)
    write_issue_report(output_dir / ISSUE_REPORT, issues, pairs)

    shutil.rmtree(temp_dir, ignore_errors=True)
    cleanup_old_temp_dirs(output_dir)

    print(f"paired_units={sum(1 for pair in pairs.values() if pair.question_pdf and pair.answer_pdf)}")
    print(f"questions={len(questions)}")
    print(f"issues={len(issues)}")
    print(f"downloaded_from_cirn={len(downloaded)}")
    print(f"manifest={target_paths_file if target_unit_keys else 'ALL_DOWNLOADS'}")
    print(f"output_dir={output_dir}")
    print(f"reading_docx={output_dir / READING_DOCX}")
    print(f"csv={output_dir / OUTPUT_CSV}")
    print(f"issues_report={output_dir / ISSUE_REPORT}")


if __name__ == "__main__":
    main()
