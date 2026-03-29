from __future__ import annotations

import csv
import hashlib
import html
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import fitz
import requests
import urllib3
from docx import Document


urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


ROOT_ESCAPED = "\\u570b\\u4e2d\\u984c\\u5eab"
CHINESE_DIR_ESCAPED = "\\u570b\\u6587"
DOWNLOAD_PREFIX_ESCAPED = "\\u570b\\u6587\\u7b2c"
QUESTION_SUFFIX_ESCAPED = "_\\u984c\\u76ee\\u5377"
ANSWER_SUFFIX_ESCAPED = "_\\u7b54\\u6848\\u5377"
NO_DATA_ESCAPED = "\\u7121\\u8cc7\\u6599\\uff01"
SEARCH_ESCAPED = "\\u641c\\u5c0b"

CIRN_BASE_URL = "https://cirn.moe.edu.tw"
CIRN_VOLUME_URLS = {
    "一": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13080",
    "二": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13081",
    "三": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13082",
    "四": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13083",
    "五": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13084",
    "六": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13085",
}

QUESTION_START_RE = re.compile(r"^[（(]\s*[）)]\s*(\d+)\s*[、．.]?\s*(.*)$")
ANSWER_START_RE = re.compile(r"^[（(]\s*([A-Fa-f])\s*[）)]\s*(\d+)\s*[、．.]?\s*(.*)$")
ROW_RE = re.compile(
    r'<span id="ctl00_ContentPlaceHolder1_gv1_ctl(?P<row>\d+)_lbFileName">(?P<name>[^<]+)</span>.*?'
    r'javascript:__doPostBack\(&#39;(?P<target>[^&#]+)&#39;,&#39;&#39;\)',
    re.S,
)


def escaped_name(path: Path) -> str:
    return path.name.encode("unicode_escape").decode()


def find_child_by_escaped(parent: Path, escaped: str) -> Path:
    return next(child for child in parent.iterdir() if escaped_name(child) == escaped)


def normalize_pdf_name(path: Path) -> str:
    escaped = escaped_name(path)
    escaped = escaped.replace(" (1).pdf", ".pdf")
    return escaped


@dataclass
class SourcePair:
    unit_key: str
    question_pdf: Path | None = None
    answer_pdf: Path | None = None


@dataclass
class ParsedQuestion:
    volume: str
    category: str
    title: str
    unit_key: str
    question_no: int
    question_text: str
    options: list[str]
    answer: str
    explanation: str
    question_file: str
    answer_file: str


def get_desktop_root() -> Path:
    return Path.home() / "Desktop"


def get_question_bank_root() -> Path:
    return find_child_by_escaped(get_desktop_root(), ROOT_ESCAPED)


def ensure_output_dir() -> Path:
    root = get_question_bank_root()
    target = root / find_display_name(CHINESE_DIR_ESCAPED)
    target.mkdir(exist_ok=True)
    return target


def find_display_name(escaped: str) -> str:
    return escaped.encode("utf-8").decode("unicode_escape")


def iter_local_chinese_pdfs(downloads_dir: Path) -> Iterable[Path]:
    for path in sorted(downloads_dir.iterdir()):
        if path.suffix.lower() != ".pdf":
            continue
        escaped = escaped_name(path)
        if escaped.startswith(DOWNLOAD_PREFIX_ESCAPED):
            yield path


def volume_from_unit_key(unit_key: str) -> str:
    volume_map = {
        "\\u4e00": "一",
        "\\u4e8c": "二",
        "\\u4e09": "三",
        "\\u56db": "四",
        "\\u4e94": "五",
        "\\u516d": "六",
    }
    for escaped, display in volume_map.items():
        prefix = f"\\u570b\\u6587\\u7b2c{escaped}\\u518a_"
        if unit_key.startswith(prefix):
            return display
    raise ValueError(f"Cannot parse volume from {unit_key}")


def unit_key_from_escaped_filename(escaped_filename: str) -> tuple[str, str | None]:
    if QUESTION_SUFFIX_ESCAPED in escaped_filename:
        return escaped_filename.split(QUESTION_SUFFIX_ESCAPED)[0], "question"
    if ANSWER_SUFFIX_ESCAPED in escaped_filename:
        return escaped_filename.split(ANSWER_SUFFIX_ESCAPED)[0], "answer"
    return escaped_filename.rsplit(".", 1)[0], None


def collect_pairs(paths: Iterable[Path]) -> dict[str, SourcePair]:
    pairs: dict[str, SourcePair] = {}
    for path in paths:
        escaped = normalize_pdf_name(path)
        unit_key, kind = unit_key_from_escaped_filename(escaped)
        pair = pairs.setdefault(unit_key, SourcePair(unit_key=unit_key))
        if kind == "question":
            current = pair.question_pdf
            if current is None or " (1).pdf" in current.name:
                pair.question_pdf = path
        elif kind == "answer":
            current = pair.answer_pdf
            if current is None or " (1).pdf" in current.name:
                pair.answer_pdf = path
    return pairs


def fetch_all_rows_for_volume(url: str) -> tuple[requests.Session, str]:
    session = requests.Session()
    page = session.get(url, timeout=30, verify=False).text
    fields = build_hidden_field_map(page)
    fields["ctl00$ContentPlaceHolder1$PagingControl1$txtPageCut"] = "200"
    fields["ctl00$ContentPlaceHolder1$PagingControl1$btnNoUsed"] = ""
    page = session.post(url, data=fields, timeout=30, verify=False).text
    return session, page


def build_hidden_field_map(page: str) -> dict[str, str]:
    return {
        match.group(1): html.unescape(match.group(2))
        for match in re.finditer(
            r'<input[^>]+type="hidden"[^>]+name="([^"]+)"[^>]+value="([^"]*)"',
            page,
        )
    }


def download_missing_from_cirn(pairs: dict[str, SourcePair], temp_dir: Path) -> list[str]:
    downloaded: list[str] = []
    for volume, url in CIRN_VOLUME_URLS.items():
        volume_units = [
            pair
            for unit_key, pair in pairs.items()
            if volume_from_unit_key(unit_key) == volume
            and (pair.question_pdf is None or pair.answer_pdf is None)
        ]
        if not volume_units:
            continue

        session, page = fetch_all_rows_for_volume(url)
        available_rows = {
            html.unescape(match.group("name")): html.unescape(match.group("target"))
            for match in ROW_RE.finditer(page)
        }

        for pair in volume_units:
            if pair.question_pdf is None:
                escaped_filename = f"{pair.unit_key}{QUESTION_SUFFIX_ESCAPED}.pdf"
                display_filename = escaped_filename.encode("utf-8").decode("unicode_escape")
                target = available_rows.get(display_filename)
                if target:
                    pair.question_pdf = download_file_from_cirn(
                        session=session,
                        page=page,
                        url=url,
                        event_target=target,
                        output_path=temp_dir / display_filename,
                    )
                    downloaded.append(display_filename)

            if pair.answer_pdf is None:
                escaped_filename = f"{pair.unit_key}{ANSWER_SUFFIX_ESCAPED}.pdf"
                display_filename = escaped_filename.encode("utf-8").decode("unicode_escape")
                target = available_rows.get(display_filename)
                if target:
                    pair.answer_pdf = download_file_from_cirn(
                        session=session,
                        page=page,
                        url=url,
                        event_target=target,
                        output_path=temp_dir / display_filename,
                    )
                    downloaded.append(display_filename)
    return downloaded


def download_file_from_cirn(
    session: requests.Session,
    page: str,
    url: str,
    event_target: str,
    output_path: Path,
) -> Path:
    fields = build_hidden_field_map(page)
    fields["ctl00$ContentPlaceHolder1$PagingControl1$txtPageCut"] = "200"
    fields["__EVENTTARGET"] = event_target
    fields["__EVENTARGUMENT"] = ""
    response = session.post(url, data=fields, timeout=60, verify=False)
    output_path.write_bytes(response.content)
    return output_path


def read_pdf_text(pdf_path: Path) -> str:
    document = fitz.open(pdf_path)
    chunks: list[str] = []
    for page in document:
        chunks.append(page.get_text("text"))
    return "\n".join(chunks)


def normalize_line(line: str) -> str:
    return re.sub(r"\s+", " ", line.replace("\xa0", " ")).strip()


def is_header_line(line: str) -> bool:
    if not line:
        return True
    if line in {"題目卷", "答案卷"}:
        return True
    if line.startswith("國文／第") or line.startswith("國文/第"):
        return True
    if line.startswith("國文 / 第"):
        return True
    if re.fullmatch(r"\d+", line):
        return True
    return False


def preprocess_pdf_lines(text: str) -> list[str]:
    lines = [normalize_line(line) for line in text.splitlines()]
    return [line for line in lines if not is_header_line(line)]


def split_question_blocks(lines: list[str]) -> dict[int, list[str]]:
    blocks: dict[int, list[str]] = {}
    current_no: int | None = None
    current_lines: list[str] = []
    for line in lines:
        match = QUESTION_START_RE.match(line)
        if match:
            if current_no is not None:
                blocks[current_no] = current_lines
            current_no = int(match.group(1))
            current_lines = []
            if match.group(2).strip():
                current_lines.append(match.group(2).strip())
            continue
        if current_no is not None:
            current_lines.append(line)
    if current_no is not None:
        blocks[current_no] = current_lines
    return blocks


def split_answer_blocks(lines: list[str]) -> dict[int, tuple[str, list[str]]]:
    blocks: dict[int, tuple[str, list[str]]] = {}
    current_no: int | None = None
    current_answer = ""
    current_lines: list[str] = []
    for line in lines:
        match = ANSWER_START_RE.match(line)
        if match:
            if current_no is not None:
                blocks[current_no] = (current_answer, current_lines)
            current_answer = match.group(1).upper()
            current_no = int(match.group(2))
            current_lines = []
            if match.group(3).strip():
                current_lines.append(match.group(3).strip())
            continue
        if current_no is not None:
            current_lines.append(line)
    if current_no is not None:
        blocks[current_no] = (current_answer, current_lines)
    return blocks


def parse_block_to_question_and_options(lines: list[str]) -> tuple[str, list[str]]:
    cleaned = [normalize_line(line) for line in lines if normalize_line(line)]
    if not cleaned:
        return "", []
    if len(cleaned) >= 5:
        options = cleaned[-4:]
        stem_lines = cleaned[:-4]
    elif len(cleaned) == 4:
        options = cleaned[-3:]
        stem_lines = cleaned[:-3]
    else:
        options = []
        stem_lines = cleaned
    question_text = "\n".join(stem_lines).strip()
    return question_text, options


def parse_unit_key(unit_key: str) -> tuple[str, str, str]:
    display = unit_key.encode("utf-8").decode("unicode_escape")
    parts = display.split("_", 2)
    if len(parts) != 3:
        raise ValueError(f"Unexpected unit key: {unit_key}")
    return parts[0], parts[1], parts[2]


def parse_pair(pair: SourcePair) -> list[ParsedQuestion]:
    if pair.question_pdf is None or pair.answer_pdf is None:
        return []

    volume, category, title = parse_unit_key(pair.unit_key)
    question_lines = preprocess_pdf_lines(read_pdf_text(pair.question_pdf))
    answer_lines = preprocess_pdf_lines(read_pdf_text(pair.answer_pdf))

    question_blocks = split_question_blocks(question_lines)
    answer_blocks = split_answer_blocks(answer_lines)
    results: list[ParsedQuestion] = []

    for question_no in sorted(question_blocks):
        block_lines = question_blocks[question_no]
        question_text, options = parse_block_to_question_and_options(block_lines)
        answer, answer_block_lines = answer_blocks.get(question_no, ("", []))
        if not question_text and answer_block_lines:
            question_text, _ = parse_block_to_question_and_options(answer_block_lines)
        results.append(
            ParsedQuestion(
                volume=volume,
                category=category,
                title=title,
                unit_key=pair.unit_key,
                question_no=question_no,
                question_text=question_text,
                options=options,
                answer=answer,
                explanation="",
                question_file=pair.question_pdf.name,
                answer_file=pair.answer_pdf.name,
            )
        )
    return results


def write_reading_txt(output_path: Path, questions: list[ParsedQuestion], missing_units: list[str]) -> None:
    with output_path.open("w", encoding="utf-8") as fh:
        fh.write("=== 國文 題庫 ===\n")
        fh.write(f"總題數: {len(questions)}\n\n")
        fh.write("=== 缺件單元 ===\n")
        if missing_units:
            for unit in missing_units:
                fh.write(f"{unit}\n")
        else:
            fh.write("無\n")
        fh.write("\n")

        for index, question in enumerate(questions, start=1):
            fh.write(
                f"第 {index} 題. [{question.volume} / {question.category} / {question.title} / 原題號 {question.question_no}]\n"
            )
            fh.write(f"題目: {question.question_text}\n")
            for label, option in zip(["A", "B", "C", "D", "E", "F"], question.options):
                fh.write(f"({label}) {option}\n")
            fh.write(f"答案: {question.answer}\n")
            if question.explanation:
                fh.write(f"解析: {question.explanation}\n")
            fh.write("----------------------------------------\n")


def write_reading_docx(output_path: Path, questions: list[ParsedQuestion], missing_units: list[str]) -> None:
    document = Document()
    document.add_heading("國文題庫閱讀版", level=0)
    document.add_paragraph(f"總題數：{len(questions)}")
    document.add_heading("缺件單元", level=1)
    if missing_units:
        for unit in missing_units:
            document.add_paragraph(unit, style="List Bullet")
    else:
        document.add_paragraph("無")

    current_unit = None
    for question in questions:
        unit_title = f"{question.volume} / {question.category} / {question.title}"
        if unit_title != current_unit:
            document.add_heading(unit_title, level=1)
            current_unit = unit_title
        document.add_paragraph(f"題號 {question.question_no}", style="List Number")
        document.add_paragraph(question.question_text)
        for label, option in zip(["A", "B", "C", "D", "E", "F"], question.options):
            document.add_paragraph(f"({label}) {option}")
        document.add_paragraph(f"答案：{question.answer}")
    document.save(output_path)


def write_csv(output_path: Path, questions: list[ParsedQuestion]) -> None:
    with output_path.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(
            [
                "冊別",
                "類型",
                "篇名",
                "來源單元",
                "題號",
                "題目",
                "選項A",
                "選項B",
                "選項C",
                "選項D",
                "選項E",
                "選項F",
                "正確答案",
                "解析",
                "題目卷檔名",
                "答案卷檔名",
            ]
        )
        for question in questions:
            options = question.options + [""] * (6 - len(question.options))
            writer.writerow(
                [
                    question.volume,
                    question.category,
                    question.title,
                    question.unit_key.encode("utf-8").decode("unicode_escape"),
                    question.question_no,
                    question.question_text,
                    *options[:6],
                    question.answer,
                    question.explanation,
                    question.question_file,
                    question.answer_file,
                ]
            )


def write_missing_txt(output_path: Path, missing_units: list[str], downloaded_files: list[str]) -> None:
    with output_path.open("w", encoding="utf-8") as fh:
        fh.write("=== 國文 缺件清單 ===\n")
        if missing_units:
            for unit in missing_units:
                fh.write(f"{unit}\n")
        else:
            fh.write("無\n")
        fh.write("\n=== 本次從 CIRN 補件 ===\n")
        if downloaded_files:
            for name in downloaded_files:
                fh.write(f"{name}\n")
        else:
            fh.write("無\n")


def hash_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def delete_exact_duplicate_pdfs(downloads_dir: Path) -> list[str]:
    removed: list[str] = []
    seen: dict[str, tuple[Path, str]] = {}
    for path in sorted(iter_local_chinese_pdfs(downloads_dir)):
        normalized = normalize_pdf_name(path)
        file_hash = hash_file(path)
        current = seen.get(normalized)
        if current and current[1] == file_hash:
            path.unlink()
            removed.append(path.name)
            continue
        if normalized not in seen or " (1).pdf" in seen[normalized][0].name:
            seen[normalized] = (path, file_hash)
    return removed


def cleanup_temp_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)


def main() -> None:
    desktop = get_desktop_root()
    downloads_dir = Path.home() / "Downloads"
    output_dir = ensure_output_dir()
    temp_dir = output_dir / "_tmp_cirn_downloads"
    temp_dir.mkdir(exist_ok=True)

    local_pairs = collect_pairs(iter_local_chinese_pdfs(downloads_dir))
    downloaded_files = download_missing_from_cirn(local_pairs, temp_dir)

    all_paths = list(iter_local_chinese_pdfs(downloads_dir))
    if temp_dir.exists():
        all_paths.extend(sorted(temp_dir.glob("*.pdf")))

    pairs = collect_pairs(all_paths)
    missing_units: list[str] = []
    for unit_key, pair in sorted(pairs.items()):
        display_unit = unit_key.encode("utf-8").decode("unicode_escape")
        if pair.question_pdf is None and pair.answer_pdf is None:
            continue
        if pair.question_pdf is None:
            missing_units.append(f"[只有答案卷] {display_unit}")
        elif pair.answer_pdf is None:
            missing_units.append(f"[只有題目卷] {display_unit}")

    questions: list[ParsedQuestion] = []
    for unit_key in sorted(pairs):
        questions.extend(parse_pair(pairs[unit_key]))

    questions.sort(
        key=lambda q: (
            q.volume,
            q.category,
            q.title,
            q.question_no,
        )
    )
    excluded_without_answer = sum(1 for question in questions if not question.answer.strip())
    excluded_incomplete_options = sum(1 for question in questions if len(question.options) < 4)
    questions = [
        question
        for question in questions
        if question.answer.strip() and len(question.options) >= 4
    ]

    write_reading_txt(output_dir / "國文_閱讀版.txt", questions, missing_units)
    write_reading_docx(output_dir / "國文_閱讀版.docx", questions, missing_units)
    write_csv(output_dir / "國文_所有題目.csv", questions)
    write_missing_txt(output_dir / "國文_缺件清單.txt", missing_units, downloaded_files)

    duplicate_deleted = delete_exact_duplicate_pdfs(downloads_dir)
    cleanup_temp_dir(temp_dir)

    print(f"paired_units={sum(1 for pair in pairs.values() if pair.question_pdf and pair.answer_pdf)}")
    print(f"questions={len(questions)}")
    print(f"missing_units={len(missing_units)}")
    print(f"downloaded_from_cirn={len(downloaded_files)}")
    print(f"deleted_duplicates={len(duplicate_deleted)}")
    print(f"excluded_without_answer={excluded_without_answer}")
    print(f"excluded_incomplete_options={excluded_incomplete_options}")


if __name__ == "__main__":
    main()
