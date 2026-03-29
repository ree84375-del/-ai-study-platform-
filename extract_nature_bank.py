from __future__ import annotations

import csv
import html
import re
import shutil
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from pathlib import Path

import requests
import urllib3
from docx import Document

from extract_math_bank import (
    ANSWER_SUFFIX,
    CandidateParts,
    QUESTION_SUFFIX,
    QuestionBlock,
    build_candidate_parts,
    build_hidden_field_map,
    choose_best_parts,
    choose_preferred,
    compute_sha256,
    convert_pdf_to_docx,
    download_file_from_cirn,
    extract_explanation,
    normalize_line,
    normalize_pdf_name,
    parse_blocks,
    read_docx_lines,
    read_pdf_lines,
    render_ocr_lines,
    unit_key_from_filename,
)


urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


DOWNLOAD_PREFIX = "自然"
OUTPUT_ROOT = Path.home() / "Desktop" / "國中題庫"
LEGACY_CACHE_DIR = OUTPUT_ROOT / "自然"
TMP_DIR = Path.cwd() / "tmp" / "_tmp_nature_extract"

CIRN_BASE_URL = "https://cirn.moe.edu.tw"
CIRN_VOLUME_URLS = {
    "第一冊": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13234",
    "第二冊": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13235",
    "第三冊": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13236",
    "第四冊": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13237",
    "第五冊": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13238",
    "第六冊": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13239",
}

ROW_RE = re.compile(
    r'<span id="ctl00_ContentPlaceHolder1_gv1_ctl(?P<row>\d+)_lbFileName">(?P<name>[^<]+)</span>.*?'
    r"javascript:__doPostBack\(&#39;(?P<target>[^&#]+)&#39;,&#39;&#39;\)",
    re.S,
)

BOOKLET_ORDER = {
    "第一冊": 1,
    "第二冊": 2,
    "第三冊": 3,
    "第四冊": 4,
    "第五冊": 5,
    "第六冊": 6,
}

NATURE_SUBJECT_MAP = {
    "生物": {
        "生命的基本單位",
        "植物的營養、運輸與感應",
        "動物的營養與運輸",
        "動物的協調與恆定",
        "生物體的運輸作用",
        "生物體的營養方式",
        "生物體的協調作用",
        "生物的生殖",
        "遺傳",
        "演化與生物多樣性",
        "生物與環境",
        "地球上的生物",
    },
    "理化": {
        "認識自然與生活科技",
        "溫度與熱量",
        "物質的形態與性質",
        "物質的組成",
        "原子與分子",
        "酸鹼鹽",
        "物質的變化",
        "波動現象",
        "光",
        "聲音",
        "生活與化學",
        "有機物質",
        "運動與力",
        "摩擦力與場力",
        "重力作用",
        "電磁作用",
        "電及其應用",
        "能源與動力",
    },
    "地科": {
        "地球在太空中",
        "地球面面觀",
        "地殼的變動",
        "地貌的變遷",
        "天氣變化",
        "地球的運轉",
        "人與地球",
        "天然災害",
    },
}

SUBJECT_HEADINGS = {
    "生物": "生物題庫閱讀版",
    "理化": "理化題庫閱讀版",
    "地科": "地科題庫閱讀版",
}


@dataclass
class SourcePair:
    unit_key: str
    question_candidates: list[Path] = field(default_factory=list)
    answer_candidates: list[Path] = field(default_factory=list)
    question_pdf: Path | None = None
    answer_pdf: Path | None = None


@dataclass
class NatureQuestion:
    subject: str
    booklet: str
    chapter: str
    topic: str
    source_unit: str
    question_number: int
    question_text: str
    option_a: str
    option_b: str
    option_c: str
    option_d: str
    correct_answer: str
    explanation: str


@dataclass
class IssueItem:
    kind: str
    subject: str
    unit_key: str
    question_number: int | None
    detail: str


def get_downloads_dir() -> Path:
    return Path.home() / "Downloads"


def collect_input_paths(downloads_dir: Path) -> list[Path]:
    return [
        path
        for path in downloads_dir.iterdir()
        if path.suffix.lower() == ".pdf" and path.name.startswith("自然第")
    ]


def collect_pairs(paths: list[Path]) -> dict[str, SourcePair]:
    pairs: dict[str, SourcePair] = {}
    for path in sorted(paths, key=lambda item: item.name):
        normalized_name = normalize_pdf_name(path.name)
        unit_key, kind = unit_key_from_filename(normalized_name)
        pair = pairs.setdefault(unit_key, SourcePair(unit_key=unit_key))
        if kind == "question":
            pair.question_candidates.append(path)
        elif kind == "answer":
            pair.answer_candidates.append(path)

    for pair in pairs.values():
        pair.question_pdf = choose_preferred(pair.question_candidates)
        pair.answer_pdf = choose_preferred(pair.answer_candidates)
    return pairs


def parse_unit_parts(unit_key: str) -> tuple[str, str, str]:
    parts = unit_key.split("_")
    booklet = parts[0] if parts else ""
    if booklet.startswith("自然"):
        booklet = booklet.replace("自然", "", 1)
    chapter = parts[1] if len(parts) > 1 else ""
    topic = "_".join(parts[2:]) if len(parts) > 2 else chapter
    return booklet, chapter, topic


def detect_subject_from_chapter(chapter: str) -> str:
    for subject, chapter_names in NATURE_SUBJECT_MAP.items():
        if chapter in chapter_names:
            return subject
    return ""


def volume_from_unit_key(unit_key: str) -> str:
    booklet, _, _ = parse_unit_parts(unit_key)
    if booklet not in CIRN_VOLUME_URLS:
        raise ValueError(f"Cannot parse volume from unit key: {unit_key}")
    return booklet


def fetch_all_rows_for_volume(url: str) -> tuple[requests.Session, str]:
    session = requests.Session()
    page = session.get(url, timeout=30, verify=False).text
    fields = build_hidden_field_map(page)
    fields["ctl00$ContentPlaceHolder1$PagingControl1$txtPageCut"] = "500"
    fields["ctl00$ContentPlaceHolder1$PagingControl1$btnNoUsed"] = ""
    page = session.post(url, data=fields, timeout=30, verify=False).text
    return session, page


def download_missing_from_cirn(pairs: dict[str, SourcePair], temp_dir: Path) -> list[str]:
    downloaded: list[str] = []

    for volume, url in CIRN_VOLUME_URLS.items():
        session, page = fetch_all_rows_for_volume(url)
        available_rows = {
            html.unescape(match.group("name")): html.unescape(match.group("target"))
            for match in ROW_RE.finditer(page)
        }

        volume_units_map: dict[str, SourcePair] = {}
        for filename in sorted(available_rows):
            if not filename.startswith(f"自然{volume}") or not filename.endswith(".pdf"):
                continue
            normalized_name = normalize_pdf_name(filename)
            unit_key, kind = unit_key_from_filename(normalized_name)
            if not kind:
                continue
            pair = pairs.setdefault(unit_key, SourcePair(unit_key=unit_key))
            volume_units_map[unit_key] = pair

        volume_units = [
            pair
            for pair in volume_units_map.values()
            if pair.question_pdf is None or pair.answer_pdf is None
        ]
        if not volume_units:
            continue

        for pair in volume_units:
            if pair.question_pdf is None:
                filename = f"{pair.unit_key}{QUESTION_SUFFIX}.pdf"
                event_target = available_rows.get(filename)
                if event_target:
                    pair.question_pdf = download_file_from_cirn(
                        session=session,
                        page=page,
                        url=url,
                        event_target=event_target,
                        output_path=temp_dir / filename,
                    )
                    pair.question_candidates.append(pair.question_pdf)
                    downloaded.append(filename)

            if pair.answer_pdf is None:
                filename = f"{pair.unit_key}{ANSWER_SUFFIX}.pdf"
                event_target = available_rows.get(filename)
                if event_target:
                    pair.answer_pdf = download_file_from_cirn(
                        session=session,
                        page=page,
                        url=url,
                        event_target=event_target,
                        output_path=temp_dir / filename,
                    )
                    pair.answer_candidates.append(pair.answer_pdf)
                    downloaded.append(filename)

    return downloaded


def get_docx_cache_path(pdf_path: Path, temp_dir: Path) -> Path:
    LEGACY_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    cached_path = LEGACY_CACHE_DIR / f"{pdf_path.stem}.docx"
    if cached_path.exists():
        return cached_path
    temp_path = temp_dir / "docx" / f"{pdf_path.stem}.docx"
    temp_path.parent.mkdir(parents=True, exist_ok=True)
    return temp_path


def load_docx_lines(pdf_path: Path, temp_dir: Path, issues: list[IssueItem], subject: str, unit_key: str, role_label: str) -> list[str]:
    docx_path = get_docx_cache_path(pdf_path, temp_dir)
    try:
        if not docx_path.exists():
            convert_pdf_to_docx(pdf_path, docx_path)
        return read_docx_lines(docx_path)
    except Exception as exc:
        issues.append(IssueItem("docx_conversion_failed", subject, unit_key, None, f"{role_label}: {exc}"))
        return []


def build_questions_for_pair(pair: SourcePair, temp_dir: Path, issues: list[IssueItem]) -> list[NatureQuestion]:
    if pair.question_pdf is None or pair.answer_pdf is None:
        return []

    booklet, chapter, topic = parse_unit_parts(pair.unit_key)
    subject = detect_subject_from_chapter(chapter)
    if not subject:
        issues.append(IssueItem("unknown_subject", "未分類", pair.unit_key, None, f"無法從章節判斷科目：{chapter}"))
        return []

    if len(pair.question_candidates) > 1:
        issues.append(
            IssueItem(
                "duplicate_download",
                subject,
                pair.unit_key,
                None,
                "題目卷有重複下載：" + "、".join(path.name for path in sorted(pair.question_candidates, key=lambda item: item.name)),
            )
        )
    if len(pair.answer_candidates) > 1:
        issues.append(
            IssueItem(
                "duplicate_download",
                subject,
                pair.unit_key,
                None,
                "答案卷有重複下載：" + "、".join(path.name for path in sorted(pair.answer_candidates, key=lambda item: item.name)),
            )
        )

    question_direct_lines = read_pdf_lines(pair.question_pdf)
    answer_direct_lines = read_pdf_lines(pair.answer_pdf)
    question_docx_lines = load_docx_lines(pair.question_pdf, temp_dir, issues, subject, pair.unit_key, "question_pdf")
    answer_docx_lines = load_docx_lines(pair.answer_pdf, temp_dir, issues, subject, pair.unit_key, "answer_pdf")

    question_blocks_direct = parse_blocks(question_direct_lines, with_answers=False)
    answer_blocks_direct = parse_blocks(answer_direct_lines, with_answers=True)
    question_blocks_docx = parse_blocks(question_docx_lines, with_answers=False)
    answer_blocks_docx = parse_blocks(answer_docx_lines, with_answers=True)

    image_dir = temp_dir / "ocr_images"
    image_dir.mkdir(parents=True, exist_ok=True)

    seed_numbers = set(question_blocks_direct) | set(question_blocks_docx) | set(answer_blocks_direct) | set(answer_blocks_docx)
    low_quality_seed_count = 0
    for number in sorted(seed_numbers):
        seed_candidates: list[CandidateParts] = []
        for source_name, block in [
            ("question_direct", question_blocks_direct.get(number)),
            ("question_docx", question_blocks_docx.get(number)),
            ("answer_direct", answer_blocks_direct.get(number)),
            ("answer_docx", answer_blocks_docx.get(number)),
        ]:
            if not block:
                continue
            candidate = build_candidate_parts(block.lines, source_name=source_name)
            if candidate:
                seed_candidates.append(candidate)
        best_seed = choose_best_parts(seed_candidates)
        if best_seed is None or len(best_seed.question_text) < 8 or len([opt for opt in best_seed.options if opt]) < 4:
            low_quality_seed_count += 1

    low_quality_threshold = max(3, len(seed_numbers) // 5) if seed_numbers else 1
    need_question_ocr = (
        not question_blocks_direct
        or len(question_blocks_direct) < len(answer_blocks_direct)
        or low_quality_seed_count >= low_quality_threshold
    )
    need_answer_ocr = (
        not answer_blocks_direct
        or len(answer_blocks_direct) < len(question_blocks_direct)
    )

    if need_question_ocr:
        try:
            question_ocr_lines = render_ocr_lines(pair.question_pdf, image_dir=image_dir)
        except Exception as exc:
            question_ocr_lines = []
            issues.append(IssueItem("ocr_failed", subject, pair.unit_key, None, f"題目卷 OCR 失敗：{exc}"))
    else:
        question_ocr_lines = []

    if need_answer_ocr:
        try:
            answer_ocr_lines = render_ocr_lines(pair.answer_pdf, image_dir=image_dir)
        except Exception as exc:
            answer_ocr_lines = []
            issues.append(IssueItem("ocr_failed", subject, pair.unit_key, None, f"答案卷 OCR 失敗：{exc}"))
    else:
        answer_ocr_lines = []

    question_blocks_ocr = parse_blocks(question_ocr_lines, with_answers=False)
    answer_blocks_ocr = parse_blocks(answer_ocr_lines, with_answers=True)

    question_source_counts = {
        "pdf": len(question_blocks_direct),
        "word": len(question_blocks_docx),
        "ocr": len(question_blocks_ocr),
    }
    answer_source_counts = {
        "pdf": len(answer_blocks_direct),
        "word": len(answer_blocks_docx),
        "ocr": len(answer_blocks_ocr),
    }
    if len({count for count in question_source_counts.values() if count}) > 1:
        issues.append(
            IssueItem(
                "question_source_count_mismatch",
                subject,
                pair.unit_key,
                None,
                f"題目卷三輪檢查題數不一致：{question_source_counts}",
            )
        )
    if len({count for count in answer_source_counts.values() if count}) > 1:
        issues.append(
            IssueItem(
                "answer_source_count_mismatch",
                subject,
                pair.unit_key,
                None,
                f"答案卷三輪檢查題數不一致：{answer_source_counts}",
            )
        )

    question_numbers = set(question_blocks_direct) | set(question_blocks_docx) | set(question_blocks_ocr)
    answer_numbers = set(answer_blocks_direct) | set(answer_blocks_docx) | set(answer_blocks_ocr)

    for number in sorted(answer_numbers - question_numbers):
        issues.append(IssueItem("answer_without_question", subject, pair.unit_key, number, "答案卷有題號，但題目卷三輪都沒抓到。"))
    for number in sorted(question_numbers - answer_numbers):
        issues.append(IssueItem("question_without_answer", subject, pair.unit_key, number, "題目卷有題號，但答案卷三輪都沒抓到。"))

    questions: list[NatureQuestion] = []
    for number in sorted(question_numbers | answer_numbers):
        block_candidates: list[CandidateParts] = []
        for source_name, block in [
            ("question_direct", question_blocks_direct.get(number)),
            ("question_docx", question_blocks_docx.get(number)),
            ("question_ocr", question_blocks_ocr.get(number)),
            ("answer_direct", answer_blocks_direct.get(number)),
            ("answer_docx", answer_blocks_docx.get(number)),
            ("answer_ocr", answer_blocks_ocr.get(number)),
        ]:
            if not block:
                continue
            candidate = build_candidate_parts(block.lines, source_name=source_name)
            if candidate:
                block_candidates.append(candidate)

        best_parts = choose_best_parts(block_candidates)
        if best_parts is None:
            issues.append(IssueItem("unparsed_question", subject, pair.unit_key, number, "無法從 PDF / Word / OCR 結果重建題目。"))
            continue

        answers_found = [
            block.answer
            for block in (answer_blocks_direct.get(number), answer_blocks_docx.get(number), answer_blocks_ocr.get(number))
            if block and block.answer
        ]
        answer = Counter(answers_found).most_common(1)[0][0] if answers_found else ""
        if len(set(answers_found)) > 1:
            issues.append(
                IssueItem(
                    "answer_mismatch",
                    subject,
                    pair.unit_key,
                    number,
                    f"PDF / Word / OCR 出現不同答案：{', '.join(sorted(set(answers_found)))}",
                )
            )

        option_count = len([option for option in best_parts.options if option])
        if len(best_parts.question_text) < 8 or option_count < 4 or answer not in {"A", "B", "C", "D"}:
            issues.append(
                IssueItem(
                    "low_quality_question",
                    subject,
                    pair.unit_key,
                    number,
                    f"question_len={len(best_parts.question_text)} option_count={option_count} answer={answer or 'EMPTY'} source={best_parts.source_name}",
                )
            )

        answer_block = answer_blocks_direct.get(number) or answer_blocks_docx.get(number) or answer_blocks_ocr.get(number)
        explanation_lines = answer_block.lines if answer_block else []
        explanation = extract_explanation(explanation_lines, answer=answer)
        if not explanation:
            explanation = f"正確答案為 {answer or '未提供'}。這題目前沒有原始詳解，建議先對照題幹關鍵字、四個選項與所屬主題「{topic}」再複習。"

        options = (best_parts.options + ["", "", "", ""])[:4]
        questions.append(
            NatureQuestion(
                subject=subject,
                booklet=booklet,
                chapter=chapter,
                topic=topic,
                source_unit=pair.unit_key,
                question_number=number,
                question_text=best_parts.question_text,
                option_a=options[0],
                option_b=options[1],
                option_c=options[2],
                option_d=options[3],
                correct_answer=answer,
                explanation=explanation,
            )
        )

    return questions


def normalize_question_text(text: str) -> str:
    return " ".join((text or "").replace("\u3000", " ").replace("\xa0", " ").split())


def deduplicate_questions(questions: list[NatureQuestion], issues: list[IssueItem], subject: str) -> list[NatureQuestion]:
    kept: list[NatureQuestion] = []
    seen: dict[str, NatureQuestion] = {}
    for question in questions:
        key = normalize_question_text(question.question_text)
        if not key:
            issues.append(IssueItem("empty_question", subject, question.source_unit, question.question_number, "題幹為空，已略過。"))
            continue
        if key in seen:
            original = seen[key]
            issues.append(
                IssueItem(
                    "duplicate_question_text",
                    subject,
                    question.source_unit,
                    question.question_number,
                    f"與 {original.source_unit} 第 {original.question_number} 題題幹完全相同，已移除重複題。",
                )
            )
            continue
        seen[key] = question
        kept.append(question)
    return kept


def filter_clean_questions(questions: list[NatureQuestion], issues: list[IssueItem], subject: str) -> list[NatureQuestion]:
    cleaned: list[NatureQuestion] = []
    for question in questions:
        options = [question.option_a, question.option_b, question.option_c, question.option_d]
        if not normalize_question_text(question.question_text):
            issues.append(IssueItem("empty_question", subject, question.source_unit, question.question_number, "題幹為空。"))
            continue
        if len([option for option in options if normalize_question_text(option)]) < 4:
            issues.append(IssueItem("incomplete_options", subject, question.source_unit, question.question_number, "選項不足四個。"))
            continue
        if question.correct_answer not in {"A", "B", "C", "D"}:
            issues.append(IssueItem("invalid_answer", subject, question.source_unit, question.question_number, "正確答案缺失或不合法。"))
            continue
        cleaned.append(question)
    return cleaned


def booklet_sort_key(booklet: str) -> int:
    return BOOKLET_ORDER.get(booklet, 99)


def get_subject_dir(subject: str) -> Path:
    subject_dir = OUTPUT_ROOT / subject
    subject_dir.mkdir(parents=True, exist_ok=True)
    return subject_dir


def write_csv(output_path: Path, questions: list[NatureQuestion]) -> None:
    fieldnames = [
        "volume",
        "category",
        "title",
        "source_unit",
        "content_text",
        "option_a",
        "option_b",
        "option_c",
        "option_d",
        "correct_answer",
        "explanation",
        "difficulty",
        "tags",
        "booklet",
        "chapter",
        "topic",
    ]
    with output_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for question in questions:
            writer.writerow(
                {
                    "volume": question.booklet,
                    "category": question.chapter,
                    "title": question.topic,
                    "source_unit": question.source_unit,
                    "content_text": question.question_text,
                    "option_a": question.option_a,
                    "option_b": question.option_b,
                    "option_c": question.option_c,
                    "option_d": question.option_d,
                    "correct_answer": question.correct_answer,
                    "explanation": question.explanation,
                    "difficulty": 2,
                    "tags": f"{question.topic} | {question.source_unit}",
                    "booklet": question.booklet,
                    "chapter": question.chapter,
                    "topic": question.topic,
                }
            )


def write_reading_txt(output_path: Path, subject: str, questions: list[NatureQuestion], missing_units: list[str], issue_count: int) -> None:
    lines = [
        f"{subject}題庫閱讀版",
        "=" * 24,
        f"總題數：{len(questions)}",
        f"缺件數：{len(missing_units)}",
        f"問題清單筆數：{issue_count}",
        "",
        "缺件單元",
        "-" * 16,
    ]
    if missing_units:
        lines.extend(f"- {unit}" for unit in missing_units)
    else:
        lines.append("無")

    current_marker: tuple[str, str, str] | None = None
    for index, question in enumerate(questions, start=1):
        marker = (question.booklet, question.chapter, question.topic)
        if marker != current_marker:
            lines.extend(["", f"{question.booklet} / {question.chapter} / {question.topic}", "-" * 40])
            current_marker = marker

        lines.extend(
            [
                f"第 {index} 題",
                question.question_text,
                f"A. {question.option_a}",
                f"B. {question.option_b}",
                f"C. {question.option_c}",
                f"D. {question.option_d}",
                f"答案：{question.correct_answer}",
                f"詳解：{question.explanation}",
                "",
            ]
        )

    output_path.write_text("\n".join(lines), encoding="utf-8")


def write_reading_docx(output_path: Path, subject: str, questions: list[NatureQuestion], missing_units: list[str], issue_count: int) -> None:
    document = Document()
    document.add_heading(SUBJECT_HEADINGS[subject], level=0)
    document.add_paragraph(f"總題數：{len(questions)}")
    document.add_paragraph(f"缺件數：{len(missing_units)}")
    document.add_paragraph(f"問題清單筆數：{issue_count}")
    document.add_heading("缺件單元", level=1)
    if missing_units:
        for unit in missing_units:
            document.add_paragraph(unit, style="List Bullet")
    else:
        document.add_paragraph("無")

    current_marker: tuple[str, str, str] | None = None
    for index, question in enumerate(questions, start=1):
        marker = (question.booklet, question.chapter, question.topic)
        if marker != current_marker:
            document.add_heading(f"{question.booklet} / {question.chapter} / {question.topic}", level=1)
            current_marker = marker

        document.add_paragraph(f"第 {index} 題", style="List Number")
        document.add_paragraph(question.question_text)
        document.add_paragraph(f"(A) {question.option_a}")
        document.add_paragraph(f"(B) {question.option_b}")
        document.add_paragraph(f"(C) {question.option_c}")
        document.add_paragraph(f"(D) {question.option_d}")
        document.add_paragraph(f"答案：{question.correct_answer}")
        document.add_paragraph(f"詳解：{question.explanation}")

    document.save(output_path)


def write_missing_report(output_path: Path, subject: str, pairs: dict[str, SourcePair]) -> None:
    lines = [f"{subject}題庫缺件清單", "=" * 24, ""]
    missing_rows: list[str] = []
    for pair in sorted(pairs.values(), key=lambda item: item.unit_key):
        _, chapter, _ = parse_unit_parts(pair.unit_key)
        if detect_subject_from_chapter(chapter) != subject:
            continue
        if pair.question_pdf is None:
            missing_rows.append(f"- {pair.unit_key}: 缺題目卷")
        if pair.answer_pdf is None:
            missing_rows.append(f"- {pair.unit_key}: 缺答案卷")

    if missing_rows:
        lines.extend(missing_rows)
    else:
        lines.append("無缺件")

    output_path.write_text("\n".join(lines), encoding="utf-8")


def write_issue_report(output_path: Path, subject: str, issues: list[IssueItem], pairs: dict[str, SourcePair]) -> None:
    sections: list[str] = [f"{subject}題庫問題題目清單", "=" * 24, ""]

    duplicate_rows = []
    for pair in sorted(pairs.values(), key=lambda item: item.unit_key):
        _, chapter, _ = parse_unit_parts(pair.unit_key)
        if detect_subject_from_chapter(chapter) != subject:
            continue
        for kind, candidates in (("題目卷", pair.question_candidates), ("答案卷", pair.answer_candidates)):
            if len(candidates) <= 1:
                continue
            hashes = {candidate.name: compute_sha256(candidate)[:12] for candidate in candidates}
            duplicate_rows.append((pair.unit_key, kind, hashes))

    sections.append("[重複下載]")
    if not duplicate_rows:
        sections.append("- 無")
    else:
        for unit_key, kind, hashes in duplicate_rows:
            sections.append(f"- {unit_key} / {kind}")
            for filename, short_hash in hashes.items():
                sections.append(f"  {filename} / sha256={short_hash}")
    sections.append("")

    grouped: dict[str, list[IssueItem]] = {}
    for issue in issues:
        if issue.subject != subject:
            continue
        grouped.setdefault(issue.kind, []).append(issue)

    for kind in sorted(grouped):
        sections.append(f"[{kind}]")
        for issue in grouped[kind]:
            if issue.question_number is None:
                sections.append(f"- {issue.unit_key}: {issue.detail}")
            else:
                sections.append(f"- {issue.unit_key} / 第 {issue.question_number} 題: {issue.detail}")
        sections.append("")

    output_path.write_text("\n".join(sections), encoding="utf-8")


def write_summary(summary_path: Path, pair_count: int, pdf_count: int, downloaded: list[str], questions_by_subject: dict[str, list[NatureQuestion]], issues: list[IssueItem]) -> None:
    lines = [
        "自然拆科整理摘要",
        "=" * 24,
        f"掃描 PDF：{pdf_count}",
        f"單元配對：{pair_count}",
        f"CIRN 補下載：{len(downloaded)}",
        "",
    ]
    for subject in ("生物", "理化", "地科"):
        lines.append(f"{subject}：{len(questions_by_subject.get(subject, []))} 題")

    lines.extend(["", "問題清單統計", "-" * 16])
    issue_counter = Counter(issue.kind for issue in issues)
    if issue_counter:
        for kind, count in sorted(issue_counter.items()):
            lines.append(f"- {kind}: {count}")
    else:
        lines.append("無")

    lines.extend(["", "CIRN 補下載檔案", "-" * 16])
    if downloaded:
        lines.extend(f"- {name}" for name in downloaded)
    else:
        lines.append("無")

    summary_path.parent.mkdir(parents=True, exist_ok=True)
    summary_path.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    downloads_dir = get_downloads_dir()
    TMP_DIR.mkdir(parents=True, exist_ok=True)

    pdf_paths = collect_input_paths(downloads_dir)
    pairs = collect_pairs(pdf_paths)
    downloaded = download_missing_from_cirn(pairs, TMP_DIR / "downloads")

    issues: list[IssueItem] = []
    raw_questions: list[NatureQuestion] = []

    for pair in sorted(pairs.values(), key=lambda item: item.unit_key):
        if pair.question_pdf is None or pair.answer_pdf is None:
            _, chapter, _ = parse_unit_parts(pair.unit_key)
            subject = detect_subject_from_chapter(chapter) or "未分類"
            if pair.question_pdf is None:
                issues.append(IssueItem("missing_question_pdf", subject, pair.unit_key, None, "缺題目卷。"))
            if pair.answer_pdf is None:
                issues.append(IssueItem("missing_answer_pdf", subject, pair.unit_key, None, "缺答案卷。"))
            continue
        raw_questions.extend(build_questions_for_pair(pair, TMP_DIR, issues))

    questions_by_subject: dict[str, list[NatureQuestion]] = defaultdict(list)
    for question in raw_questions:
        questions_by_subject[question.subject].append(question)

    for subject in ("生物", "理化", "地科"):
        subject_questions = sorted(
            questions_by_subject.get(subject, []),
            key=lambda item: (booklet_sort_key(item.booklet), item.chapter, item.topic, item.question_number),
        )
        subject_questions = deduplicate_questions(subject_questions, issues, subject)
        subject_questions = filter_clean_questions(subject_questions, issues, subject)
        subject_questions = sorted(
            subject_questions,
            key=lambda item: (booklet_sort_key(item.booklet), item.chapter, item.topic, item.question_number),
        )
        questions_by_subject[subject] = subject_questions

        subject_dir = get_subject_dir(subject)
        missing_units = [
            pair.unit_key
            for pair in sorted(pairs.values(), key=lambda item: item.unit_key)
            if detect_subject_from_chapter(parse_unit_parts(pair.unit_key)[1]) == subject
            and (pair.question_pdf is None or pair.answer_pdf is None)
        ]
        write_csv(subject_dir / f"{subject}_所有題目.csv", subject_questions)
        write_reading_txt(subject_dir / f"{subject}_閱讀版.txt", subject, subject_questions, missing_units, len([item for item in issues if item.subject == subject]))
        write_reading_docx(subject_dir / f"{subject}_閱讀版.docx", subject, subject_questions, missing_units, len([item for item in issues if item.subject == subject]))
        write_missing_report(subject_dir / f"{subject}_缺件清單.txt", subject, pairs)
        write_issue_report(subject_dir / f"{subject}_問題題目清單.txt", subject, issues, pairs)

    write_summary(
        OUTPUT_ROOT / "自然" / "自然_整理摘要.txt",
        pair_count=len(pairs),
        pdf_count=len(pdf_paths),
        downloaded=downloaded,
        questions_by_subject=questions_by_subject,
        issues=issues,
    )

    shutil.rmtree(TMP_DIR, ignore_errors=True)
    print(f"paired_units={sum(1 for pair in pairs.values() if pair.question_pdf and pair.answer_pdf)}")
    print(f"biology={len(questions_by_subject.get('生物', []))}")
    print(f"integrated_science={len(questions_by_subject.get('理化', []))}")
    print(f"earth_science={len(questions_by_subject.get('地科', []))}")
    print(f"downloaded_from_cirn={len(downloaded)}")
    print(f"summary={OUTPUT_ROOT / '自然' / '自然_整理摘要.txt'}")


if __name__ == "__main__":
    main()
