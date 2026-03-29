from __future__ import annotations

import csv
import hashlib
import html
import logging
import re
import shutil
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from pathlib import Path

import fitz
import requests
import urllib3
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from pdf2docx import Converter
from rapidocr_onnxruntime import RapidOCR


urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
logging.getLogger("pdf2docx").setLevel(logging.ERROR)
logging.getLogger("pdf2docx.converter").setLevel(logging.ERROR)


DOWNLOAD_PREFIX = "社會"
QUESTION_SUFFIX = "_題目卷"
ANSWER_SUFFIX = "_答案卷"
OUTPUT_ROOT = Path.home() / "Desktop" / "國中題庫"
TMP_DIR = Path.cwd() / "tmp" / "_tmp_social_extract"
MAIN_CIRN_URL = "https://cirn.moe.edu.tw/WebFile/index.aspx?sid=1193&mid=13209"

ROW_RE = re.compile(
    r'<span id="ctl00_ContentPlaceHolder1_gv1_ctl(?P<row>\d+)_lbFileName">(?P<name>[^<]+)</span>.*?'
    r"javascript:__doPostBack\(&#39;(?P<target>[^&#]+)&#39;,&#39;&#39;\)",
    re.S,
)

QUESTION_START_PATTERNS = [
    re.compile(r"^[（(]\s*[）)]\s*(\d+)\s*[、．.:：]?\s*(.*)$"),
    re.compile(r"^問題\s*(\d+)\s*[、．.:：]?\s*(.*)$"),
]
ANSWER_START_PATTERNS = [
    re.compile(r"^[（(]\s*([A-D])\s*[）)]\s*(\d+)\s*[、．.:：]?\s*(.*)$", re.I),
    re.compile(r"^([A-D])\s*[、．.:：)]\s*(\d+)\s*[、．.:：]?\s*(.*)$", re.I),
]
OPTION_START_PATTERNS = [
    re.compile(r"^[（(]?\s*([A-D])\s*[）).:：]\s*(.*)$", re.I),
]
SUBITEM_RE = re.compile(r"^[（(]?\s*\d+\s*[）)]")
OCR_ENGINE = RapidOCR()

BOOKLET_ORDER = {
    "第一冊": 1,
    "第二冊": 2,
    "第三冊": 3,
    "第四冊": 4,
    "第五冊": 5,
    "第六冊": 6,
}

SOCIAL_SUBJECT_MAP = {
    "地理": {
        "臺灣的自然環境",
        "臺灣人文環境",
        "中國的自然人文環境",
        "中國大地風情畫及全球概論",
        "世界風情畫",
        "世界風情畫-非洲與大洋洲",
        "全球關聯",
    },
    "歷史": {
        "臺灣的歷史",
        "臺灣歷史發展",
        "中國的歷史",
        "近代中國的劇變",
        "世界文明的發展",
        "近現代世界歷史",
    },
    "公民": {
        "個人成長與群體",
        "社會生活與文化",
        "民主政治與生活",
        "法律與生活",
        "經濟與生活",
        "全球化與世界公民",
    },
}

SUBJECT_ICON_HEADINGS = {
    "地理": "地理題庫閱讀版",
    "歷史": "歷史題庫閱讀版",
    "公民": "公民題庫閱讀版",
}


@dataclass
class SourcePair:
    unit_key: str
    question_candidates: list[Path] = field(default_factory=list)
    answer_candidates: list[Path] = field(default_factory=list)
    question_pdf: Path | None = None
    answer_pdf: Path | None = None


@dataclass
class QuestionBlock:
    number: int
    answer: str = ""
    lines: list[str] = field(default_factory=list)


@dataclass
class CandidateParts:
    question_text: str
    options: list[str]
    source_name: str
    score: int


@dataclass
class SocialQuestion:
    subject: str
    booklet: str
    chapter: str
    topic: str
    source_unit: str
    question_number: int
    question_text: str
    option_a: str
    option_b: str
    option_c: str
    option_d: str
    correct_answer: str
    explanation: str


@dataclass
class IssueItem:
    kind: str
    subject: str
    unit_key: str
    question_number: int | None
    detail: str


def get_downloads_dir() -> Path:
    return Path.home() / "Downloads"


def normalize_pdf_name(name: str) -> str:
    cleaned = re.sub(rf"{re.escape(QUESTION_SUFFIX)}\s*\(\d+\)(?=\.pdf$)", QUESTION_SUFFIX, name)
    cleaned = re.sub(rf"{re.escape(ANSWER_SUFFIX)}\s*\(\d+\)(?=\.pdf$)", ANSWER_SUFFIX, cleaned)
    cleaned = re.sub(r"\s*\(\d+\)(?=\.pdf$)", "", cleaned)
    return cleaned


def unit_key_from_filename(name: str) -> tuple[str, str | None]:
    if QUESTION_SUFFIX in name:
        return name.split(QUESTION_SUFFIX)[0], "question"
    if ANSWER_SUFFIX in name:
        return name.split(ANSWER_SUFFIX)[0], "answer"
    return name.rsplit(".", 1)[0], None


def duplicate_penalty(path: Path) -> tuple[int, int, str]:
    normalized = normalize_pdf_name(path.name)
    return (
        1 if path.name != normalized else 0,
        len(path.name),
        path.name,
    )


def choose_preferred(paths: list[Path]) -> Path | None:
    if not paths:
        return None
    return sorted(paths, key=duplicate_penalty)[0]


def collect_input_paths(downloads_dir: Path) -> list[Path]:
    return [
        path
        for path in downloads_dir.iterdir()
        if path.suffix.lower() == ".pdf" and path.name.startswith(DOWNLOAD_PREFIX)
    ]


def collect_pairs(paths: list[Path]) -> dict[str, SourcePair]:
    pairs: dict[str, SourcePair] = {}
    for path in sorted(paths, key=lambda item: item.name):
        normalized_name = normalize_pdf_name(path.name)
        unit_key, kind = unit_key_from_filename(normalized_name)
        pair = pairs.setdefault(unit_key, SourcePair(unit_key=unit_key))
        if kind == "question":
            pair.question_candidates.append(path)
        elif kind == "answer":
            pair.answer_candidates.append(path)

    for pair in pairs.values():
        pair.question_pdf = choose_preferred(pair.question_candidates)
        pair.answer_pdf = choose_preferred(pair.answer_candidates)
    return pairs


def build_hidden_field_map(page: str) -> dict[str, str]:
    return {
        match.group(1): html.unescape(match.group(2))
        for match in re.finditer(
            r'<input[^>]+type="hidden"[^>]+name="([^"]+)"[^>]+value="([^"]*)"',
            page,
        )
    }


def fetch_all_rows() -> tuple[requests.Session, str]:
    session = requests.Session()
    page = session.get(MAIN_CIRN_URL, timeout=30, verify=False).text
    fields = build_hidden_field_map(page)
    fields["ctl00$ContentPlaceHolder1$PagingControl1$txtPageCut"] = "500"
    fields["ctl00$ContentPlaceHolder1$PagingControl1$btnNoUsed"] = ""
    page = session.post(MAIN_CIRN_URL, data=fields, timeout=30, verify=False).text
    return session, page


def download_file_from_cirn(
    session: requests.Session,
    page: str,
    event_target: str,
    output_path: Path,
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    fields = build_hidden_field_map(page)
    fields["ctl00$ContentPlaceHolder1$PagingControl1$txtPageCut"] = "500"
    fields["__EVENTTARGET"] = event_target
    fields["__EVENTARGUMENT"] = ""
    response = session.post(MAIN_CIRN_URL, data=fields, timeout=60, verify=False)
    response.raise_for_status()
    output_path.write_bytes(response.content)
    return output_path


def download_missing_from_cirn(pairs: dict[str, SourcePair], temp_dir: Path) -> list[str]:
    missing_pairs = [pair for pair in pairs.values() if pair.question_pdf is None or pair.answer_pdf is None]
    if not missing_pairs:
        return []

    session, page = fetch_all_rows()
    available_rows = {
        html.unescape(match.group("name")): html.unescape(match.group("target"))
        for match in ROW_RE.finditer(page)
    }
    downloaded: list[str] = []

    for pair in missing_pairs:
        if pair.question_pdf is None:
            filename = f"{pair.unit_key}{QUESTION_SUFFIX}.pdf"
            target = available_rows.get(filename)
            if target:
                downloaded_path = download_file_from_cirn(session, page, target, temp_dir / filename)
                pair.question_candidates.append(downloaded_path)
                pair.question_pdf = downloaded_path
                downloaded.append(filename)

        if pair.answer_pdf is None:
            filename = f"{pair.unit_key}{ANSWER_SUFFIX}.pdf"
            target = available_rows.get(filename)
            if target:
                downloaded_path = download_file_from_cirn(session, page, target, temp_dir / filename)
                pair.answer_candidates.append(downloaded_path)
                pair.answer_pdf = downloaded_path
                downloaded.append(filename)

    return downloaded


def normalize_line(line: str) -> str:
    line = (
        line.replace("\xa0", " ")
        .replace("\u3000", " ")
        .replace("\u200b", "")
        .replace("\ufeff", "")
    )
    return re.sub(r"\s+", " ", line).strip()


def is_header_line(line: str) -> bool:
    if not line:
        return True
    if line.startswith("社會／第") or line.startswith("社會/第"):
        return True
    if line in {"【題目卷】", "【答案卷】"}:
        return True
    if line.startswith("頁") and re.search(r"\d+$", line):
        return True
    if re.fullmatch(r"\d+", line):
        return True
    return False


def read_pdf_lines(pdf_path: Path) -> list[str]:
    document = fitz.open(pdf_path)
    lines: list[str] = []
    for page in document:
        for raw_line in page.get_text("text").splitlines():
            line = normalize_line(raw_line)
            if not line or is_header_line(line):
                continue
            lines.append(line)
    document.close()
    return lines


def convert_pdf_to_docx(pdf_path: Path, output_path: Path) -> None:
    converter = Converter(str(pdf_path))
    try:
        converter.convert(str(output_path), start=0, end=None)
    finally:
        converter.close()


def iter_block_items(parent: DocumentObject | _Cell):
    parent_element = parent.element.body if isinstance(parent, DocumentObject) else parent._tc
    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def read_docx_lines(docx_path: Path) -> list[str]:
    document = Document(docx_path)
    lines: list[str] = []

    def append_line(value: str) -> None:
        line = normalize_line(value)
        if line and not is_header_line(line):
            lines.append(line)

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            append_line(block.text)
            continue
        for row in block.rows:
            for cell in row.cells:
                for item in iter_block_items(cell):
                    if isinstance(item, Paragraph):
                        append_line(item.text)
    return lines


def render_ocr_lines(pdf_path: Path, image_dir: Path) -> list[str]:
    document = fitz.open(pdf_path)
    lines: list[str] = []
    for index, page in enumerate(document):
        image_path = image_dir / f"{pdf_path.stem}_page_{index + 1}.png"
        pixmap = page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
        pixmap.save(image_path)
        ocr_result, _ = OCR_ENGINE(str(image_path))
        for item in ocr_result or []:
            line = normalize_line(str(item[1]))
            if line and not is_header_line(line):
                lines.append(line)
    document.close()
    return lines


def parse_blocks(lines: list[str], with_answers: bool) -> dict[int, QuestionBlock]:
    patterns = ANSWER_START_PATTERNS if with_answers else QUESTION_START_PATTERNS
    blocks: dict[int, QuestionBlock] = {}
    current: QuestionBlock | None = None
    preface_lines: list[str] = []

    def should_buffer_for_next_question(current_lines: list[str], incoming_line: str) -> bool:
        if with_answers:
            return False
        labeled_count = 0
        for existing_line in current_lines:
            if any(pattern.match(existing_line) for pattern in OPTION_START_PATTERNS):
                labeled_count += 1
        if labeled_count >= 4 and len(incoming_line) > 18 and not SUBITEM_RE.match(incoming_line):
            return True

        trailing_short = 0
        for existing_line in reversed(current_lines):
            if any(pattern.match(existing_line) for pattern in OPTION_START_PATTERNS):
                break
            if len(existing_line) <= 120:
                trailing_short += 1
            else:
                break
        if trailing_short >= 4 and len(incoming_line) > 18:
            return True
        return False

    for raw_line in lines:
        line = normalize_line(raw_line)
        if not line or is_header_line(line):
            continue

        matched = False
        for pattern in patterns:
            match = pattern.match(line)
            if not match:
                continue

            if current is not None:
                blocks[current.number] = current

            if with_answers:
                answer, number, rest = match.groups()
                current = QuestionBlock(number=int(number), answer=answer.upper())
            else:
                number, rest = match.groups()
                current = QuestionBlock(number=int(number))

            if preface_lines:
                current.lines.extend(preface_lines)
                preface_lines = []

            rest = normalize_line(rest)
            if rest:
                current.lines.append(rest)
            matched = True
            break

        if matched:
            continue

        if current is not None:
            if should_buffer_for_next_question(current.lines, line):
                preface_lines.append(line)
                continue
            current.lines.append(line)
        else:
            preface_lines.append(line)

    if current is not None:
        blocks[current.number] = current

    return blocks


def parse_labeled_options(lines: list[str]) -> tuple[list[str], list[str]]:
    question_lines: list[str] = []
    options_map: dict[str, list[str]] = {}
    current_option: str | None = None

    for line in lines:
        matched_option = False
        for pattern in OPTION_START_PATTERNS:
            match = pattern.match(line)
            if not match:
                continue
            label, rest = match.groups()
            label = label.upper()
            if label not in {"A", "B", "C", "D"}:
                continue
            options_map.setdefault(label, [])
            if rest:
                options_map[label].append(normalize_line(rest))
            current_option = label
            matched_option = True
            break

        if matched_option:
            continue

        if current_option is None:
            question_lines.append(line)
        else:
            options_map.setdefault(current_option, []).append(line)

    if set(options_map) >= {"A", "B", "C", "D"}:
        options = [normalize_line(" ".join(options_map.get(label, []))) for label in ("A", "B", "C", "D")]
        return question_lines, options
    return question_lines, []


def parse_trailing_options(lines: list[str]) -> tuple[str, list[str]]:
    if len(lines) < 5:
        return normalize_line(" ".join(lines)), []

    candidate = lines[-4:]
    if not all(line and len(line) <= 160 for line in candidate):
        return normalize_line(" ".join(lines)), []

    question_text = normalize_line(" ".join(lines[:-4]))
    options = [normalize_line(line) for line in candidate]
    return question_text, options


def score_candidate(question_text: str, options: list[str], source_name: str) -> int:
    score = min(len(question_text), 240)
    if len(options) == 4:
        score += 120
    elif len(options) == 3:
        score += 55
    else:
        score -= 80

    if source_name.endswith("_ocr"):
        score += 18
    if source_name.endswith("_direct"):
        score += 10
    if source_name.endswith("_docx"):
        score += 8

    repeated_markers = sum(question_text.count(token) for token in ("(1)", "（1）", "(A)", "（A）"))
    if repeated_markers >= 2:
        score -= 150

    if len(question_text) < 8:
        score -= 80

    if any(not option for option in options):
        score -= 40

    return score


def build_candidate_parts(lines: list[str], source_name: str) -> CandidateParts | None:
    cleaned = [normalize_line(line) for line in lines if normalize_line(line)]
    if not cleaned:
        return None

    question_lines, labeled_options = parse_labeled_options(cleaned)
    if labeled_options:
        question_text = normalize_line(" ".join(question_lines))
        return CandidateParts(
            question_text=question_text,
            options=labeled_options,
            source_name=source_name,
            score=score_candidate(question_text, labeled_options, source_name),
        )

    question_text, trailing_options = parse_trailing_options(cleaned)
    return CandidateParts(
        question_text=question_text,
        options=trailing_options,
        source_name=source_name,
        score=score_candidate(question_text, trailing_options, source_name),
    )


def choose_best_parts(candidates: list[CandidateParts]) -> CandidateParts | None:
    if not candidates:
        return None
    return max(candidates, key=lambda item: item.score)


def extract_explanation(lines: list[str], answer: str) -> str:
    for index, line in enumerate(lines):
        if "解析" in line or "詳解" in line:
            explanation_lines = [line]
            explanation_lines.extend(lines[index + 1 :])
            explanation = normalize_line(" ".join(explanation_lines))
            if explanation:
                return explanation
    if answer:
        return f"正確答案為 {answer}。原始來源未附詳解，建議先回頭比對題幹條件與正確選項。"
    return ""


def compute_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(65536), b""):
            digest.update(chunk)
    return digest.hexdigest()


def parse_unit_parts(unit_key: str) -> tuple[str, str, str]:
    parts = unit_key.split("_", 2)
    booklet = parts[0].replace("社會", "", 1) if parts else ""
    chapter = parts[1] if len(parts) > 1 else ""
    topic = parts[2] if len(parts) > 2 else chapter
    return booklet, chapter, topic


def detect_subject_from_chapter(chapter: str) -> str:
    for subject, chapter_names in SOCIAL_SUBJECT_MAP.items():
        if chapter in chapter_names:
            return subject
    return ""


def build_questions_for_pair(
    pair: SourcePair,
    temp_dir: Path,
    issues: list[IssueItem],
) -> list[SocialQuestion]:
    if pair.question_pdf is None or pair.answer_pdf is None:
        return []

    booklet, chapter, topic = parse_unit_parts(pair.unit_key)
    subject = detect_subject_from_chapter(chapter)
    if not subject:
        issues.append(IssueItem("unknown_subject", "未分類", pair.unit_key, None, f"無法從章節判斷科目：{chapter}"))
        return []

    if len(pair.question_candidates) > 1:
        issues.append(
            IssueItem(
                "duplicate_download",
                subject,
                pair.unit_key,
                None,
                "題目卷有重複下載：" + "、".join(path.name for path in sorted(pair.question_candidates, key=lambda item: item.name)),
            )
        )
    if len(pair.answer_candidates) > 1:
        issues.append(
            IssueItem(
                "duplicate_download",
                subject,
                pair.unit_key,
                None,
                "答案卷有重複下載：" + "、".join(path.name for path in sorted(pair.answer_candidates, key=lambda item: item.name)),
            )
        )

    image_dir = temp_dir / "ocr_images"
    image_dir.mkdir(parents=True, exist_ok=True)

    question_direct_lines = read_pdf_lines(pair.question_pdf)
    answer_direct_lines = read_pdf_lines(pair.answer_pdf)

    question_blocks_direct = parse_blocks(question_direct_lines, with_answers=False)
    answer_blocks_direct = parse_blocks(answer_direct_lines, with_answers=True)
    direct_seed_numbers = set(question_blocks_direct) | set(answer_blocks_direct)
    low_quality_direct_count = 0
    for number in sorted(direct_seed_numbers):
        direct_candidates: list[CandidateParts] = []
        for source_name, block in [
            ("question_direct", question_blocks_direct.get(number)),
            ("answer_direct", answer_blocks_direct.get(number)),
        ]:
            if not block:
                continue
            candidate = build_candidate_parts(block.lines, source_name=source_name)
            if candidate:
                direct_candidates.append(candidate)
        best_direct = choose_best_parts(direct_candidates)
        if best_direct is None or len(best_direct.question_text) < 8 or len([opt for opt in best_direct.options if opt]) < 4:
            low_quality_direct_count += 1

    question_blocks_docx: dict[int, QuestionBlock] = {}
    answer_blocks_docx: dict[int, QuestionBlock] = {}

    seed_numbers = set(question_blocks_direct) | set(question_blocks_docx) | set(answer_blocks_direct) | set(answer_blocks_docx)
    low_quality_seed_count = 0
    for number in sorted(seed_numbers):
        seed_candidates: list[CandidateParts] = []
        for source_name, block in [
            ("question_direct", question_blocks_direct.get(number)),
            ("question_docx", question_blocks_docx.get(number)),
            ("answer_direct", answer_blocks_direct.get(number)),
            ("answer_docx", answer_blocks_docx.get(number)),
        ]:
            if not block:
                continue
            candidate = build_candidate_parts(block.lines, source_name=source_name)
            if candidate:
                seed_candidates.append(candidate)
        best_seed = choose_best_parts(seed_candidates)
        if best_seed is None or len(best_seed.question_text) < 8 or len([opt for opt in best_seed.options if opt]) < 4:
            low_quality_seed_count += 1

    low_quality_threshold = max(3, len(direct_seed_numbers) // 5) if direct_seed_numbers else 1
    need_question_ocr = (
        not question_blocks_direct
        or len(question_blocks_direct) < len(answer_blocks_direct)
        or low_quality_seed_count >= low_quality_threshold
    )
    need_answer_ocr = (
        not answer_blocks_direct
        or len(answer_blocks_direct) < len(question_blocks_direct)
    )

    if need_question_ocr:
        try:
            question_ocr_lines = render_ocr_lines(pair.question_pdf, image_dir=image_dir)
        except Exception as exc:
            question_ocr_lines = []
            issues.append(IssueItem("ocr_failed", subject, pair.unit_key, None, f"題目卷 OCR 失敗：{exc}"))
    else:
        question_ocr_lines = []

    if need_answer_ocr:
        try:
            answer_ocr_lines = render_ocr_lines(pair.answer_pdf, image_dir=image_dir)
        except Exception as exc:
            answer_ocr_lines = []
            issues.append(IssueItem("ocr_failed", subject, pair.unit_key, None, f"答案卷 OCR 失敗：{exc}"))
    else:
        answer_ocr_lines = []

    question_blocks_ocr = parse_blocks(question_ocr_lines, with_answers=False)
    answer_blocks_ocr = parse_blocks(answer_ocr_lines, with_answers=True)

    question_source_counts = {
        "pdf": len(question_blocks_direct),
        "word": len(question_blocks_docx),
        "ocr": len(question_blocks_ocr),
    }
    answer_source_counts = {
        "pdf": len(answer_blocks_direct),
        "word": len(answer_blocks_docx),
        "ocr": len(answer_blocks_ocr),
    }
    if len({count for count in question_source_counts.values() if count}) > 1:
        issues.append(
            IssueItem(
                "question_source_count_mismatch",
                subject,
                pair.unit_key,
                None,
                f"題目卷三輪檢查題數不一致：{question_source_counts}",
            )
        )
    if len({count for count in answer_source_counts.values() if count}) > 1:
        issues.append(
            IssueItem(
                "answer_source_count_mismatch",
                subject,
                pair.unit_key,
                None,
                f"答案卷三輪檢查題數不一致：{answer_source_counts}",
            )
        )

    question_numbers = set(question_blocks_direct) | set(question_blocks_docx) | set(question_blocks_ocr)
    answer_numbers = set(answer_blocks_direct) | set(answer_blocks_docx) | set(answer_blocks_ocr)

    for number in sorted(answer_numbers - question_numbers):
        issues.append(IssueItem("answer_without_question", subject, pair.unit_key, number, "答案卷有題號，但題目卷三輪都沒抓到。"))
    for number in sorted(question_numbers - answer_numbers):
        issues.append(IssueItem("question_without_answer", subject, pair.unit_key, number, "題目卷有題號，但答案卷三輪都沒抓到。"))

    questions: list[SocialQuestion] = []
    for number in sorted(question_numbers | answer_numbers):
        block_candidates: list[CandidateParts] = []
        for source_name, block in [
            ("question_direct", question_blocks_direct.get(number)),
            ("question_docx", question_blocks_docx.get(number)),
            ("question_ocr", question_blocks_ocr.get(number)),
            ("answer_direct", answer_blocks_direct.get(number)),
            ("answer_docx", answer_blocks_docx.get(number)),
            ("answer_ocr", answer_blocks_ocr.get(number)),
        ]:
            if not block:
                continue
            candidate = build_candidate_parts(block.lines, source_name=source_name)
            if candidate:
                block_candidates.append(candidate)

        best_parts = choose_best_parts(block_candidates)
        if best_parts is None:
            issues.append(IssueItem("unparsed_question", subject, pair.unit_key, number, "PDF / Word / OCR 都無法還原題目。"))
            continue

        answers_found = [
            block.answer
            for block in (answer_blocks_direct.get(number), answer_blocks_docx.get(number), answer_blocks_ocr.get(number))
            if block and block.answer
        ]
        answer = Counter(answers_found).most_common(1)[0][0] if answers_found else ""
        if len(set(answers_found)) > 1:
            issues.append(
                IssueItem(
                    "answer_mismatch",
                    subject,
                    pair.unit_key,
                    number,
                    f"答案卷三輪檢查出現不同答案：{', '.join(sorted(set(answers_found)))}",
                )
            )

        option_count = len([option for option in best_parts.options if option])
        if len(best_parts.question_text) < 8 or option_count < 4 or answer not in {"A", "B", "C", "D"}:
            issues.append(
                IssueItem(
                    "low_quality_question",
                    subject,
                    pair.unit_key,
                    number,
                    f"question_len={len(best_parts.question_text)} option_count={option_count} answer={answer or 'EMPTY'} source={best_parts.source_name}",
                )
            )

        answer_block = answer_blocks_direct.get(number) or answer_blocks_docx.get(number) or answer_blocks_ocr.get(number)
        explanation_lines = answer_block.lines if answer_block else []
        explanation = extract_explanation(explanation_lines, answer=answer)

        options = (best_parts.options + ["", "", "", ""])[:4]
        questions.append(
            SocialQuestion(
                subject=subject,
                booklet=booklet,
                chapter=chapter,
                topic=topic,
                source_unit=pair.unit_key,
                question_number=number,
                question_text=best_parts.question_text,
                option_a=options[0],
                option_b=options[1],
                option_c=options[2],
                option_d=options[3],
                correct_answer=answer,
                explanation=explanation,
            )
        )

    return questions


def normalize_question_text(text: str) -> str:
    return " ".join((text or "").replace("\u3000", " ").replace("\xa0", " ").split())


def deduplicate_questions(questions: list[SocialQuestion], issues: list[IssueItem]) -> list[SocialQuestion]:
    kept: list[SocialQuestion] = []
    seen: dict[str, SocialQuestion] = {}
    for question in questions:
        key = normalize_question_text(question.question_text)
        if not key:
            issues.append(IssueItem("empty_question", question.subject, question.source_unit, question.question_number, "題幹為空，已略過。"))
            continue
        if key in seen:
            original = seen[key]
            issues.append(
                IssueItem(
                    "duplicate_question_text",
                    question.subject,
                    question.source_unit,
                    question.question_number,
                    f"與 {original.source_unit} 第 {original.question_number} 題題幹完全相同，已移除重複題。",
                )
            )
            continue
        seen[key] = question
        kept.append(question)
    return kept


def filter_clean_questions(questions: list[SocialQuestion], issues: list[IssueItem]) -> list[SocialQuestion]:
    cleaned: list[SocialQuestion] = []
    for question in questions:
        options = [question.option_a, question.option_b, question.option_c, question.option_d]
        if not normalize_question_text(question.question_text):
            issues.append(IssueItem("empty_question", question.subject, question.source_unit, question.question_number, "題幹為空。"))
            continue
        if len([option for option in options if normalize_question_text(option)]) < 4:
            issues.append(IssueItem("incomplete_options", question.subject, question.source_unit, question.question_number, "選項不足四個。"))
            continue
        if question.correct_answer not in {"A", "B", "C", "D"}:
            issues.append(IssueItem("invalid_answer", question.subject, question.source_unit, question.question_number, "正確答案缺失或不合法。"))
            continue
        cleaned.append(question)
    return cleaned


def booklet_sort_key(booklet: str) -> int:
    return BOOKLET_ORDER.get(booklet, 99)


def write_csv(output_path: Path, questions: list[SocialQuestion]) -> None:
    fieldnames = [
        "volume",
        "category",
        "title",
        "source_unit",
        "content_text",
        "option_a",
        "option_b",
        "option_c",
        "option_d",
        "correct_answer",
        "explanation",
        "difficulty",
        "tags",
        "booklet",
        "chapter",
        "topic",
    ]
    with output_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for question in questions:
            writer.writerow(
                {
                    "volume": question.booklet,
                    "category": question.chapter,
                    "title": question.topic,
                    "source_unit": question.source_unit,
                    "content_text": question.question_text,
                    "option_a": question.option_a,
                    "option_b": question.option_b,
                    "option_c": question.option_c,
                    "option_d": question.option_d,
                    "correct_answer": question.correct_answer,
                    "explanation": question.explanation,
                    "difficulty": 2,
                    "tags": f"{question.topic} | {question.source_unit}",
                    "booklet": question.booklet,
                    "chapter": question.chapter,
                    "topic": question.topic,
                }
            )


def write_reading_txt(output_path: Path, subject: str, questions: list[SocialQuestion], missing_units: list[str], issue_count: int) -> None:
    lines = [
        f"{subject}題庫閱讀版",
        "=" * 24,
        f"總題數：{len(questions)}",
        f"缺件數：{len(missing_units)}",
        f"問題清單筆數：{issue_count}",
        "",
        "缺件單元",
        "-" * 16,
    ]
    if missing_units:
        lines.extend(f"- {unit}" for unit in missing_units)
    else:
        lines.append("無")

    current_marker: tuple[str, str, str] | None = None
    for index, question in enumerate(questions, start=1):
        marker = (question.booklet, question.chapter, question.topic)
        if marker != current_marker:
            lines.extend(["", f"{question.booklet} / {question.chapter} / {question.topic}", "-" * 40])
            current_marker = marker

        lines.extend(
            [
                f"第 {index} 題",
                f"原單元題號：{question.question_number}",
                question.question_text,
                f"A. {question.option_a}",
                f"B. {question.option_b}",
                f"C. {question.option_c}",
                f"D. {question.option_d}",
                f"答案：{question.correct_answer}",
                f"詳解：{question.explanation}",
                "",
            ]
        )

    output_path.write_text("\n".join(lines), encoding="utf-8")


def write_reading_docx(output_path: Path, subject: str, questions: list[SocialQuestion], missing_units: list[str], issue_count: int) -> None:
    document = Document()
    document.add_heading(SUBJECT_ICON_HEADINGS[subject], level=0)
    document.add_paragraph(f"總題數：{len(questions)}")
    document.add_paragraph(f"缺件數：{len(missing_units)}")
    document.add_paragraph(f"問題清單筆數：{issue_count}")
    document.add_heading("缺件單元", level=1)
    if missing_units:
        for unit in missing_units:
            document.add_paragraph(unit, style="List Bullet")
    else:
        document.add_paragraph("無")

    current_marker: tuple[str, str, str] | None = None
    for index, question in enumerate(questions, start=1):
        marker = (question.booklet, question.chapter, question.topic)
        if marker != current_marker:
            document.add_heading(f"{question.booklet} / {question.chapter} / {question.topic}", level=1)
            current_marker = marker

        document.add_paragraph(f"第 {index} 題（原單元題號 {question.question_number}）", style="List Number")
        document.add_paragraph(question.question_text)
        document.add_paragraph(f"(A) {question.option_a}")
        document.add_paragraph(f"(B) {question.option_b}")
        document.add_paragraph(f"(C) {question.option_c}")
        document.add_paragraph(f"(D) {question.option_d}")
        document.add_paragraph(f"答案：{question.correct_answer}")
        document.add_paragraph(f"詳解：{question.explanation}")

    document.save(output_path)


def write_issue_report(path: Path, subject: str, issues: list[IssueItem]) -> None:
    counter = Counter(item.kind for item in issues)
    lines = [
        f"{subject}題庫問題題目清單",
        "=" * 28,
        "",
        "統計",
        "-" * 10,
    ]
    if counter:
        for kind, count in sorted(counter.items()):
            lines.append(f"- {kind}: {count}")
    else:
        lines.append("無")

    lines.extend(["", "明細", "-" * 10])
    if not issues:
        lines.append("無")
    else:
        for item in issues:
            qno = f"第 {item.question_number} 題" if item.question_number is not None else "整份單元"
            lines.append(f"- [{item.kind}] {item.unit_key} / {qno} / {item.detail}")

    path.write_text("\n".join(lines), encoding="utf-8")


def write_missing_report(path: Path, subject: str, pairs: dict[str, SourcePair]) -> None:
    lines = [f"{subject}題庫缺件清單", "=" * 24, ""]
    missing_rows = []
    for pair in sorted(pairs.values(), key=lambda item: item.unit_key):
        _, chapter, _ = parse_unit_parts(pair.unit_key)
        if detect_subject_from_chapter(chapter) != subject:
            continue
        if pair.question_pdf is None:
            missing_rows.append((pair.unit_key, "缺題目卷"))
        if pair.answer_pdf is None:
            missing_rows.append((pair.unit_key, "缺答案卷"))

    if not missing_rows:
        lines.append("無缺件")
    else:
        for unit_key, detail in missing_rows:
            lines.append(f"- {unit_key}: {detail}")

    path.write_text("\n".join(lines), encoding="utf-8")


def clean_existing_outputs(subject: str) -> Path:
    subject_dir = OUTPUT_ROOT / subject
    subject_dir.mkdir(parents=True, exist_ok=True)
    for path in subject_dir.iterdir():
        if path.is_file():
            path.unlink()
    return subject_dir


def clean_old_social_outputs() -> None:
    social_dir = OUTPUT_ROOT / "社會"
    if not social_dir.exists():
        return
    for path in social_dir.iterdir():
        if path.is_file():
            path.unlink()


def subject_issue_partition(issues: list[IssueItem], subject: str) -> list[IssueItem]:
    return [item for item in issues if item.subject == subject]


def main() -> None:
    downloads_dir = get_downloads_dir()
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    paths = collect_input_paths(downloads_dir)
    pairs = collect_pairs(paths)
    downloaded = download_missing_from_cirn(pairs, TMP_DIR / "downloads")

    issues: list[IssueItem] = []
    raw_questions: list[SocialQuestion] = []

    for pair in sorted(pairs.values(), key=lambda item: item.unit_key):
        if pair.question_pdf is None or pair.answer_pdf is None:
            _, chapter, _ = parse_unit_parts(pair.unit_key)
            subject = detect_subject_from_chapter(chapter) or "未分類"
            if pair.question_pdf is None:
                issues.append(IssueItem("missing_question_pdf", subject, pair.unit_key, None, "缺題目卷。"))
            if pair.answer_pdf is None:
                issues.append(IssueItem("missing_answer_pdf", subject, pair.unit_key, None, "缺答案卷。"))
            continue
        raw_questions.extend(build_questions_for_pair(pair, TMP_DIR, issues))

    questions_by_subject: dict[str, list[SocialQuestion]] = defaultdict(list)
    for question in raw_questions:
        questions_by_subject[question.subject].append(question)

    summary_lines = ["社會題庫整理摘要", "=" * 24]
    summary_lines.append(f"掃描 PDF：{len(paths)}")
    summary_lines.append(f"單元配對：{len(pairs)}")
    summary_lines.append(f"CIRN 補下載：{len(downloaded)}")

    for subject in ("地理", "歷史", "公民"):
        subject_questions = sorted(
            questions_by_subject.get(subject, []),
            key=lambda item: (booklet_sort_key(item.booklet), item.chapter, item.topic, item.question_number),
        )
        subject_questions = deduplicate_questions(subject_questions, issues)
        subject_questions = filter_clean_questions(subject_questions, issues)
        subject_questions = sorted(
            subject_questions,
            key=lambda item: (booklet_sort_key(item.booklet), item.chapter, item.topic, item.question_number),
        )
        subject_dir = clean_existing_outputs(subject)
        subject_issues = subject_issue_partition(issues, subject)
        missing_units = [
            pair.unit_key
            for pair in sorted(pairs.values(), key=lambda item: item.unit_key)
            if detect_subject_from_chapter(parse_unit_parts(pair.unit_key)[1]) == subject
            and (pair.question_pdf is None or pair.answer_pdf is None)
        ]

        write_csv(subject_dir / f"{subject}_所有題目.csv", subject_questions)
        write_reading_txt(subject_dir / f"{subject}_閱讀版.txt", subject, subject_questions, missing_units, len(subject_issues))
        write_reading_docx(subject_dir / f"{subject}_閱讀版.docx", subject, subject_questions, missing_units, len(subject_issues))
        write_missing_report(subject_dir / f"{subject}_缺件清單.txt", subject, pairs)
        write_issue_report(subject_dir / f"{subject}_問題題目清單.txt", subject, subject_issues)

        summary_lines.append(f"{subject}：{len(subject_questions)} 題")

    clean_old_social_outputs()
    summary_lines.extend(["", "CIRN 補下載檔案", "-" * 16])
    if downloaded:
        summary_lines.extend(f"- {name}" for name in downloaded)
    else:
        summary_lines.append("無")

    social_dir = OUTPUT_ROOT / "社會"
    social_dir.mkdir(parents=True, exist_ok=True)
    (social_dir / "社會_整理摘要.txt").write_text("\n".join(summary_lines), encoding="utf-8")

    print("\n".join(summary_lines))
    shutil.rmtree(TMP_DIR, ignore_errors=True)


if __name__ == "__main__":
    main()
