from __future__ import annotations

import csv
import html
import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

import fitz
import requests
import urllib3
from docx import Document


urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


CIRN_BASE_URL = "https://cirn.moe.edu.tw"
CIRN_VOLUME_URLS = {
    "一": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13208",
    "二": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13209",
    "三": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13210",
    "四": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13211",
    "五": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13212",
    "六": f"{CIRN_BASE_URL}/WebFile/index.aspx?sid=1193&mid=13213",
}
QUESTION_SUFFIX = "_題目卷"
ANSWER_SUFFIX = "_答案卷"
ROW_RE = re.compile(
    r'<span id="ctl00_ContentPlaceHolder1_gv1_ctl(?P<row>\d+)_lbFileName">(?P<name>[^<]+)</span>.*?'
    r'javascript:__doPostBack\(&#39;(?P<target>[^&#]+)&#39;,&#39;&#39;\)',
    re.S,
)
QUESTION_START_RE = re.compile(r"^[（(]\s*[）)]\s*(\d+)\s*[、.．]?\s*(.*)$")
ANSWER_START_RE = re.compile(r"^[（(]\s*([A-Fa-f])\s*[）)]\s*(\d+)\s*[、.．]?\s*(.*)$")


@dataclass
class SourcePair:
    unit_key: str
    question_pdf: Path | None = None
    answer_pdf: Path | None = None


@dataclass
class QuestionBlock:
    question_no: int
    answer: str = ""
    lines: list[str] = field(default_factory=list)


@dataclass
class ParsedQuestion:
    volume: str
    category: str
    title: str
    unit_key: str
    question_no: int
    question_text: str
    options: list[str]
    answer: str
    explanation: str
    question_file: str
    answer_file: str


def get_downloads_dir() -> Path:
    return Path.home() / "Downloads"


def get_output_dir() -> Path:
    output_dir = Path.home() / "Desktop" / "國中題庫" / "英文"
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def iter_local_english_pdfs(downloads_dir: Path) -> Iterable[Path]:
    for path in sorted(downloads_dir.iterdir()):
        if path.suffix.lower() != ".pdf":
            continue
        if path.name.startswith("英語第"):
            yield path


def normalize_pdf_name(path: Path) -> str:
    return re.sub(r" \(\d+\)(?=\.pdf$)", "", path.name)


def unit_key_from_filename(filename: str) -> tuple[str, str | None]:
    if QUESTION_SUFFIX in filename:
        return filename.split(QUESTION_SUFFIX)[0], "question"
    if ANSWER_SUFFIX in filename:
        return filename.split(ANSWER_SUFFIX)[0], "answer"
    return filename.rsplit(".", 1)[0], None


def choose_preferred(existing: Path | None, candidate: Path) -> Path:
    if existing is None:
        return candidate
    existing_penalty = 1 if re.search(r" \(\d+\)\.pdf$", existing.name) else 0
    candidate_penalty = 1 if re.search(r" \(\d+\)\.pdf$", candidate.name) else 0
    if candidate_penalty < existing_penalty:
        return candidate
    return existing


def collect_pairs(paths: Iterable[Path]) -> dict[str, SourcePair]:
    pairs: dict[str, SourcePair] = {}
    for path in paths:
        normalized_name = normalize_pdf_name(path)
        unit_key, kind = unit_key_from_filename(normalized_name)
        pair = pairs.setdefault(unit_key, SourcePair(unit_key=unit_key))
        if kind == "question":
            pair.question_pdf = choose_preferred(pair.question_pdf, path)
        elif kind == "answer":
            pair.answer_pdf = choose_preferred(pair.answer_pdf, path)
    return pairs


def volume_from_unit_key(unit_key: str) -> str:
    match = re.match(r"英語第([一二三四五六])冊_", unit_key)
    if not match:
        raise ValueError(f"Cannot parse volume from unit key: {unit_key}")
    return match.group(1)


def build_hidden_field_map(page: str) -> dict[str, str]:
    return {
        match.group(1): html.unescape(match.group(2))
        for match in re.finditer(
            r'<input[^>]+type="hidden"[^>]+name="([^"]+)"[^>]+value="([^"]*)"',
            page,
        )
    }


def fetch_all_rows_for_volume(url: str) -> tuple[requests.Session, str]:
    session = requests.Session()
    page = session.get(url, timeout=30, verify=False).text
    fields = build_hidden_field_map(page)
    fields["ctl00$ContentPlaceHolder1$PagingControl1$txtPageCut"] = "200"
    fields["ctl00$ContentPlaceHolder1$PagingControl1$btnNoUsed"] = ""
    page = session.post(url, data=fields, timeout=30, verify=False).text
    return session, page


def download_file_from_cirn(
    session: requests.Session,
    page: str,
    url: str,
    event_target: str,
    output_path: Path,
) -> Path:
    fields = build_hidden_field_map(page)
    fields["ctl00$ContentPlaceHolder1$PagingControl1$txtPageCut"] = "200"
    fields["__EVENTTARGET"] = event_target
    fields["__EVENTARGUMENT"] = ""
    response = session.post(url, data=fields, timeout=60, verify=False)
    response.raise_for_status()
    output_path.write_bytes(response.content)
    return output_path


def download_missing_from_cirn(pairs: dict[str, SourcePair], temp_dir: Path) -> list[str]:
    downloaded: list[str] = []
    for volume, url in CIRN_VOLUME_URLS.items():
        volume_units = [
            pair
            for pair in pairs.values()
            if volume_from_unit_key(pair.unit_key) == volume
            and (pair.question_pdf is None or pair.answer_pdf is None)
        ]
        if not volume_units:
            continue

        session, page = fetch_all_rows_for_volume(url)
        available_rows = {
            html.unescape(match.group("name")): html.unescape(match.group("target"))
            for match in ROW_RE.finditer(page)
        }

        for pair in volume_units:
            if pair.question_pdf is None:
                filename = f"{pair.unit_key}{QUESTION_SUFFIX}.pdf"
                event_target = available_rows.get(filename)
                if event_target:
                    pair.question_pdf = download_file_from_cirn(
                        session=session,
                        page=page,
                        url=url,
                        event_target=event_target,
                        output_path=temp_dir / filename,
                    )
                    downloaded.append(filename)

            if pair.answer_pdf is None:
                filename = f"{pair.unit_key}{ANSWER_SUFFIX}.pdf"
                event_target = available_rows.get(filename)
                if event_target:
                    pair.answer_pdf = download_file_from_cirn(
                        session=session,
                        page=page,
                        url=url,
                        event_target=event_target,
                        output_path=temp_dir / filename,
                    )
                    downloaded.append(filename)

    return downloaded


def read_pdf_lines(pdf_path: Path) -> list[str]:
    document = fitz.open(pdf_path)
    lines: list[str] = []
    for page in document:
        for raw_line in page.get_text("text").splitlines():
            line = normalize_line(raw_line)
            if not line:
                continue
            if line.startswith("英語／第") or line.startswith("英語/第"):
                continue
            if line in {"【題目卷】", "【答案卷】"}:
                continue
            if re.fullmatch(r"\d+", line):
                continue
            lines.append(line)
    return lines


def normalize_line(line: str) -> str:
    line = line.replace("\xa0", " ").replace("\u3000", " ")
    line = line.replace("：", ":").replace("，", ",")
    return re.sub(r"\s+", " ", line).strip()


def parse_question_blocks(lines: list[str]) -> dict[int, QuestionBlock]:
    blocks: dict[int, QuestionBlock] = {}
    current: QuestionBlock | None = None

    for line in lines:
        match = QUESTION_START_RE.match(line)
        if match:
            if current is not None:
                blocks[current.question_no] = current
            current = QuestionBlock(question_no=int(match.group(1)))
            if match.group(2).strip():
                current.lines.append(match.group(2).strip())
            continue
        if current is not None:
            current.lines.append(line)

    if current is not None:
        blocks[current.question_no] = current
    return blocks


def parse_answer_blocks(lines: list[str]) -> dict[int, QuestionBlock]:
    blocks: dict[int, QuestionBlock] = {}
    current: QuestionBlock | None = None

    for line in lines:
        match = ANSWER_START_RE.match(line)
        if match:
            if current is not None:
                blocks[current.question_no] = current
            current = QuestionBlock(question_no=int(match.group(2)), answer=match.group(1).upper())
            if match.group(3).strip():
                current.lines.append(match.group(3).strip())
            continue
        if current is not None:
            current.lines.append(line)

    if current is not None:
        blocks[current.question_no] = current
    return blocks


def is_source_line(line: str) -> bool:
    return "英語輔導團" in line or (line.startswith("【") and line.endswith("】"))


def option_like(line: str) -> bool:
    if not line:
        return False
    if is_source_line(line):
        return False
    if looks_like_instruction(line):
        return False
    if QUESTION_START_RE.match(line) or ANSWER_START_RE.match(line):
        return False
    if len(line) > 160:
        return False
    return True


def looks_like_instruction(line: str) -> bool:
    return line.startswith("請閱讀以下") or line.startswith("請依據下列") or line.startswith("請根據下列")


def looks_like_separator(line: str) -> bool:
    if not line:
        return True
    if looks_like_instruction(line):
        return True
    if line.startswith("(") and line.endswith(")"):
        return True
    if line.count(":") >= 1 and len(line) < 80:
        return True
    return False


def find_inline_options(lines: list[str]) -> tuple[list[str], list[str]]:
    for start in range(1, len(lines)):
        for option_count in (4, 3):
            candidate = lines[start : start + option_count]
            if len(candidate) != option_count:
                continue
            if not all(option_like(line) for line in candidate):
                continue

            after = lines[start + option_count :]
            if after and not looks_like_separator(after[0]):
                continue

            return lines[:start], candidate
    return lines, []


def split_stem_and_options(lines: list[str]) -> tuple[str, list[str]]:
    cleaned = [line for line in lines if line]
    if not cleaned:
        return "", []

    marker_index = next((index for index, line in enumerate(cleaned) if is_source_line(line)), None)
    if marker_index is not None:
        stem_lines = cleaned[:marker_index]
        option_pool = [line for line in cleaned[marker_index + 1 :] if option_like(line)]
    else:
        stem_lines = cleaned[:]
        option_pool = []

    if len(option_pool) >= 4:
        options = option_pool[-4:]
    elif len(option_pool) >= 3:
        options = option_pool[-3:]
    else:
        fallback_pool = [line for line in cleaned if not is_source_line(line)]
        stem_lines, options = find_inline_options(fallback_pool)
        if not options:
            if len(fallback_pool) >= 4 and all(option_like(line) for line in fallback_pool[-4:]):
                options = fallback_pool[-4:]
                stem_lines = fallback_pool[:-4]
            elif len(fallback_pool) >= 3 and all(option_like(line) for line in fallback_pool[-3:]):
                options = fallback_pool[-3:]
                stem_lines = fallback_pool[:-3]
            else:
                options = []

    question_text = "\n".join(line for line in stem_lines if not is_source_line(line)).strip()
    return question_text, [option.strip() for option in options]


def parse_unit(pair: SourcePair) -> tuple[list[ParsedQuestion], list[str]]:
    issues: list[str] = []
    if pair.answer_pdf is None:
        issues.append(f"{pair.unit_key}：缺答案卷")
        return [], issues

    answer_lines = read_pdf_lines(pair.answer_pdf)
    answer_blocks = parse_answer_blocks(answer_lines)
    question_blocks = parse_question_blocks(read_pdf_lines(pair.question_pdf)) if pair.question_pdf else {}

    parts = pair.unit_key.split("_")
    volume = parts[0] if parts else pair.unit_key
    category = parts[1] if len(parts) > 1 else ""
    title = "_".join(parts[2:]) if len(parts) > 2 else pair.unit_key

    parsed_questions: list[ParsedQuestion] = []
    for question_no in sorted(answer_blocks):
        answer_block = answer_blocks[question_no]
        base_lines = answer_block.lines
        if question_no in question_blocks and question_blocks[question_no].lines:
            base_lines = question_blocks[question_no].lines

        question_text, options = split_stem_and_options(base_lines)
        if not question_text or len(options) < 2 or answer_block.answer not in {"A", "B", "C", "D"}:
            issues.append(f"{pair.unit_key}：第 {question_no} 題解析失敗")
            continue

        option_index = ord(answer_block.answer) - ord("A")
        if option_index >= len(options):
            issues.append(f"{pair.unit_key}：第 {question_no} 題答案超出選項數")
            continue

        parsed_questions.append(
            ParsedQuestion(
                volume=volume,
                category=category,
                title=title,
                unit_key=pair.unit_key,
                question_no=question_no,
                question_text=question_text,
                options=options,
                answer=answer_block.answer,
                explanation="",
                question_file=pair.question_pdf.name if pair.question_pdf else "",
                answer_file=pair.answer_pdf.name if pair.answer_pdf else "",
            )
        )

    if pair.question_pdf is None:
        issues.append(f"{pair.unit_key}：缺題目卷（已用答案卷內容整理）")

    return parsed_questions, issues


def normalize_question_text(text: str) -> str:
    return " ".join((text or "").replace("\u3000", " ").replace("\xa0", " ").split())


def dedupe_questions(questions: list[ParsedQuestion]) -> tuple[list[ParsedQuestion], int]:
    deduped: list[ParsedQuestion] = []
    seen_questions: set[str] = set()
    deleted_duplicates = 0

    for question in questions:
        normalized = normalize_question_text(question.question_text)
        if normalized in seen_questions:
            deleted_duplicates += 1
            continue
        seen_questions.add(normalized)
        deduped.append(question)

    return deduped, deleted_duplicates


def write_csv(questions: list[ParsedQuestion], output_path: Path) -> None:
    fieldnames = [
        "冊別",
        "類型",
        "篇名",
        "來源單元",
        "題號",
        "題目",
        "選項A",
        "選項B",
        "選項C",
        "選項D",
        "選項E",
        "選項F",
        "正確答案",
        "解析",
        "題目卷檔名",
        "答案卷檔名",
    ]

    with output_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for question in questions:
            option_values = question.options + ["", "", "", ""]
            writer.writerow(
                {
                    "冊別": question.volume,
                    "類型": question.category,
                    "篇名": question.title,
                    "來源單元": question.unit_key,
                    "題號": question.question_no,
                    "題目": question.question_text,
                    "選項A": option_values[0],
                    "選項B": option_values[1],
                    "選項C": option_values[2],
                    "選項D": option_values[3],
                    "選項E": "",
                    "選項F": "",
                    "正確答案": question.answer,
                    "解析": question.explanation,
                    "題目卷檔名": question.question_file,
                    "答案卷檔名": question.answer_file,
                }
            )


def write_text(questions: list[ParsedQuestion], output_path: Path) -> None:
    lines: list[str] = []
    current_unit = None
    for question in questions:
        unit_header = f"{question.volume}｜{question.category}｜{question.title}"
        if unit_header != current_unit:
            if current_unit is not None:
                lines.append("")
            lines.append(unit_header)
            lines.append("=" * len(unit_header))
            current_unit = unit_header
        lines.append(f"第 {question.question_no} 題")
        lines.append(question.question_text)
        for index, option in enumerate(question.options):
            label = chr(ord("A") + index)
            lines.append(f"{label}. {option}")
        lines.append(f"答案：{question.answer}")
        lines.append(f"解析：{question.explanation or '無解析'}")
        lines.append("")

    output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def write_docx(questions: list[ParsedQuestion], output_path: Path) -> None:
    document = Document()
    current_unit = None

    for question in questions:
        unit_header = f"{question.volume}｜{question.category}｜{question.title}"
        if unit_header != current_unit:
            document.add_heading(unit_header, level=1)
            current_unit = unit_header

        document.add_paragraph(f"第 {question.question_no} 題")
        document.add_paragraph(question.question_text)
        for index, option in enumerate(question.options):
            label = chr(ord("A") + index)
            document.add_paragraph(f"{label}. {option}")
        document.add_paragraph(f"答案：{question.answer}")
        document.add_paragraph(f"解析：{question.explanation or '無解析'}")

    document.save(output_path)


def write_missing_report(
    output_path: Path,
    missing_units: list[str],
    downloaded: list[str],
    deleted_duplicates: int,
    parsed_count: int,
) -> None:
    lines = [
        f"正式題目數：{parsed_count}",
        f"刪除重複題數：{deleted_duplicates}",
        "",
        "從 CIRN 補下載：",
    ]
    if downloaded:
        lines.extend(f"- {name}" for name in downloaded)
    else:
        lines.append("- 無")

    lines.extend(["", "仍需人工檢查 / 缺件："])
    if missing_units:
        lines.extend(f"- {item}" for item in missing_units)
    else:
        lines.append("- 無")

    output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    downloads_dir = get_downloads_dir()
    output_dir = get_output_dir()
    temp_dir = output_dir / "_tmp_cirn_downloads"

    if temp_dir.exists():
        shutil.rmtree(temp_dir)
    temp_dir.mkdir(parents=True, exist_ok=True)

    try:
        pairs = collect_pairs(iter_local_english_pdfs(downloads_dir))
        downloaded = download_missing_from_cirn(pairs, temp_dir)

        parsed_questions: list[ParsedQuestion] = []
        issues: list[str] = []
        for pair in sorted(pairs.values(), key=lambda item: item.unit_key):
            unit_questions, unit_issues = parse_unit(pair)
            parsed_questions.extend(unit_questions)
            issues.extend(unit_issues)

        deduped_questions, deleted_duplicates = dedupe_questions(parsed_questions)

        csv_path = output_dir / "英文_所有題目.csv"
        txt_path = output_dir / "英文_閱讀版.txt"
        docx_path = output_dir / "英文_閱讀版.docx"
        report_path = output_dir / "英文_缺件清單.txt"

        write_csv(deduped_questions, csv_path)
        write_text(deduped_questions, txt_path)
        write_docx(deduped_questions, docx_path)
        write_missing_report(report_path, issues, downloaded, deleted_duplicates, len(deduped_questions))

        print(f"paired_units={sum(1 for pair in pairs.values() if pair.answer_pdf is not None)}")
        print(f"questions={len(deduped_questions)}")
        print(f"issues={len(issues)}")
        print(f"downloaded_from_cirn={len(downloaded)}")
        print(f"deleted_duplicates={deleted_duplicates}")
        print(f"csv={csv_path}")
        print(f"txt={txt_path}")
        print(f"docx={docx_path}")
        print(f"report={report_path}")
    finally:
        if temp_dir.exists():
            shutil.rmtree(temp_dir)


if __name__ == "__main__":
    main()
